import os
import io
import json
import re
import PyPDF2
import docx
import requests
import pandas as pd
from datetime import datetime
from pypdf import PdfReader, errors as pypdf_errors
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import time
import logging
from hr_app.models import CareerPage

# Configure logging for better visibility
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Import docx for .docx file handling
try:
    from docx import Document
except ImportError:
    logging.warning("python-docx not installed. DOCX file processing will not be available.")
    Document = None

# OCR Imports
try:
    from PIL import Image
    from pdf2image import convert_from_bytes, PdfPageCountError, PdfSyntaxError
    import pytesseract
    # --- OCR Configuration (IMPORTANT: Adjust this if Tesseract is not in your PATH) ---
    # For Windows, you might need to specify the full path to tesseract.exe, e.g.:
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    # For Linux/macOS, if tesseract is in your PATH, this line might not be strictly necessary.
    # pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract' # Example for macOS Homebrew
    # ---------------------------------------------------------------------------------
except ImportError:
    logging.warning("OCR libraries (Pillow, pdf2image, pytesseract) not installed. Image-based PDF processing will not be available.")
    Image = None
    convert_from_bytes = None
    pytesseract = None

# Load API keys from Django settings
from django.conf import settings

BLANDAI_API_KEY = getattr(settings, 'BLANDAI_API_KEY', None)
GOOGLE_API_KEY = getattr(settings, 'GOOGLE_API_KEY', None)

# Initialize LLM model globally. It will be re-initialized in llm_call if needed.
llm = None
if GOOGLE_API_KEY:
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", google_api_key=GOOGLE_API_KEY, temperature=0.7)
        logging.info("Global LLM instance initialized successfully.")
    except Exception as e:
        logging.error(f"Failed to initialize global LLM instance: {e}")
        llm = None
else:
    logging.warning("WARNING: GOOGLE_API_KEY not set in Django settings. LLM functionality will be limited.")

# Base URL for the Gemini API (used for direct requests if not using Langchain's invoke)
GEMINI_API_BASE_URL = "[https://generativelanguage.googleapis.com/v1beta/models/](https://generativelanguage.googleapis.com/v1beta/models/)"
GEMINI_MODEL_NAME = "gemini-2.5-flash-preview-05-20"


def extract_text_from_pdf(pdf_file_obj):
    """
    Extracts text from a PDF file object using pypdf.
    If direct extraction fails, attempts OCR fallback using pdf2image and pytesseract.
    """
    extracted_text = ""
    
    # Try direct text extraction first
    try:
        reader = PdfReader(pdf_file_obj)
        if reader.is_encrypted:
            logging.warning("PDF is encrypted. Direct text extraction might fail.")
            # You might add reader.decrypt('password') here if you have a known password
        
        for page in reader.pages:
            extracted_text += page.extract_text() or ""
        
        if extracted_text.strip():
            logging.info(f"Successfully extracted {len(extracted_text)} characters directly from PDF.")
            return extracted_text.strip()
        else:
            logging.warning("Direct PDF text extraction yielded no text. Attempting OCR fallback.")

    except (pypdf_errors.PdfReadError, pypdf_errors.FileTruncatedError) as pdf_err:
        logging.warning(f"pypdf error during direct read: {pdf_err}. Falling back to OCR.")
    except Exception as e:
        logging.warning(f"Unexpected error during direct PDF extraction: {e}. Falling back to OCR.")

    # Fallback to OCR if direct extraction failed or yielded no text
    if convert_from_bytes and pytesseract and Image:
        try:
            # Reset stream pointer for convert_from_bytes
            pdf_file_obj.seek(0) 
            images = convert_from_bytes(pdf_file_obj.read(), dpi=300) # Higher DPI for better OCR
            
            if not images:
                logging.warning("pdf2image could not convert any pages to images.")
                return ""

            ocr_text_parts = []
            for i, image in enumerate(images):
                logging.debug(f"Performing OCR on page {i+1}...")
                page_ocr_text = pytesseract.image_to_string(image)
                if page_ocr_text:
                    ocr_text_parts.append(page_ocr_text.strip())
            
            extracted_text = "\n".join(ocr_text_parts)
            if extracted_text.strip():
                logging.info(f"Successfully extracted {len(extracted_text)} characters from PDF using OCR.")
                return extracted_text.strip()
            else:
                logging.warning("OCR from PDF yielded no text. It might be entirely blank or unreadable images.")

        except pytesseract.TesseractNotFoundError:
            logging.error("Tesseract OCR engine not found. Please install Tesseract and ensure its path is correctly configured.")
        except (PdfPageCountError, PdfSyntaxError) as pdf2img_err:
            logging.error(f"pdf2image error converting PDF to images: {pdf2img_err}. Check Poppler installation.")
        except Exception as e:
            logging.error(f"Unexpected error during OCR process: {e}")
    else:
        logging.warning("OCR libraries are not available. Cannot perform OCR fallback for PDF.")
        
    return "" # Return empty string if all extraction methods fail

def extract_text_from_docx(docx_file_obj):
    """
    Extracts text from a DOCX file object.
    """
    if Document is None:
        logging.error("python-docx is not installed. Cannot process .docx files.")
        return ""
    try:
        document = Document(docx_file_obj)
        text = []
        for paragraph in document.paragraphs:
            text.append(paragraph.text)
        extracted_text = "\n".join(text).strip()
        logging.info(f"Successfully extracted {len(extracted_text)} characters from DOCX.")
        return extracted_text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(txt_file_obj):
    """
    Extracts text from a plain text file object.
    """
    try:
        extracted_text = txt_file_obj.read().decode('utf-8').strip()
        logging.info(f"Successfully extracted {len(extracted_text)} characters from TXT.")
        return extracted_text
    except Exception as e:
        logging.error(f"Error extracting text from TXT: {e}")
        return ""

def extract_text_from_document(file_obj, filename):
    """
    Extracts text from various document file types based on extension.
    This function handles the file object directly, including seeking to the beginning.
    """
    if not file_obj:
        logging.warning("extract_text_from_document received no file object.")
        return ""

    ext = os.path.splitext(filename)[1].lower()
    
    # Ensure file pointer is at the beginning for reading
    file_obj.seek(0)
    
    logging.info(f"Attempting to extract text from '{filename}' (Extension: {ext}). File size: {file_obj.size} bytes.")

    try:
        if ext == '.pdf':
            return extract_text_from_pdf(file_obj)
        elif ext == '.docx':
            return extract_text_from_docx(file_obj)
        elif ext == '.txt':
            return extract_text_from_txt(file_obj)
        else:
            logging.error(f"Unsupported file type for text extraction: {ext} for file {filename}")
            return ""
    except Exception as e:
        logging.error(f"Failed to extract text from '{filename}': {e}")
        return ""


# --- HELPER FUNCTIONS FOR SALARY CALCULATION ---

def _parse_experience(experience_str):
    """Parses experience string (e.g., '5 years 3 months') into total years (float)."""
    if not experience_str or experience_str == "Not Found":
        return 0.0
    
    years = 0.0
    
    # Match years
    year_match = re.search(r'(\d+)\s+year', experience_str, re.IGNORECASE)
    if year_match:
        years += int(year_match.group(1))
        
    # Match months
    month_match = re.search(r'(\d+)\s+month', experience_str, re.IGNORECASE)
    if month_match:
        years += int(month_match.group(1)) / 12.0
        
    return years

def _calculate_salary_range(job_role, total_years_experience):
    """
    Calculates a suggested salary range (in LPA) based on job role and experience.
    """
    
    # 1. Base Salary Tiering (All values are in Lakhs Per Annum - LPA)
    
    # 🛑 CUSTOMIZE THESE LPA VALUES BASED ON YOUR SALARY BANDS
    if total_years_experience < 2:
        # Example: 2.5 LPA to 4.5 LPA
        min_lpa, max_lpa = 2.5, 4.5
    elif 2 <= total_years_experience < 5:
        # Example: 5.0 LPA to 9.0 LPA
        min_lpa, max_lpa = 5.0, 9.0
    elif 5 <= total_years_experience < 8:
        # Example: 9.0 LPA to 15.0 LPA
        min_lpa, max_lpa = 9.0, 15.0
    else: # 8+ years
        # Example: 14.0 LPA to 25.0 LPA
        min_lpa, max_lpa = 14.0, 25.0 

    # 2. Role-Based Multiplier (Example)
    role_multipliers = {
        "Software Engineer": 1.1,
        "Data Scientist": 1.2,
        "QA Tester": 0.9,
    }
    
    normalized_role = job_role.lower().strip()
    
    multiplier = 1.0
    for key, mult in role_multipliers.items():
        if key.lower() in normalized_role or normalized_role in key.lower():
            multiplier = mult
            break

    # Apply multiplier and round to one decimal place for clean LPA display
    final_min = round(min_lpa * multiplier, 1)
    final_max = round(max_lpa * multiplier, 1)

    # Returns the value in the requested LPA format
    return f"₹{final_min} LPA - ₹{final_max} LPA"


def llm_call(resume_text, job_role, experience_info, job_description_text=None):
    """
    Calls the LLM (Gemini) to analyze the resume and generate a JSON summary.
    Includes target experience criteria and optional job description in the prompt for tailored analysis.
    """
    global llm 

    if llm is None and GOOGLE_API_KEY:
        try:
            llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", google_api_key=GOOGLE_API_KEY, temperature=0.7)
            logging.info("LLM initialized successfully within llm_call.")
        except Exception as e:
            logging.error(f"Failed to initialize LLM in llm_call: {e}")
            return {"error": f"Failed to initialize LLM: {e}"}
    elif llm is None:
        logging.error("LLM not initialized. GOOGLE_API_KEY might be missing or invalid.")
        return {"error": "LLM not initialized. GOOGLE_API_KEY might be missing."}

    job_description_section = ""
    if job_description_text:
        job_description_section = f"""
**Job Description for Reference:**
{job_description_text}
"""

    # --- SIMPLIFIED PROMPT ---
    # 🛑 FIX: Removed Salary Guidelines from the prompt to avoid LLM conflicts.
    summary_template = """
**YOUR RESPONSE MUST BE A SINGLE, VALID JSON OBJECT. NO OTHER TEXT, NO MARKDOWN FENCES (```json), NO EXPLANATIONS.**

You are an expert HR evaluator. Your primary goal is to provide a **comprehensive and complete JSON analysis** of the candidate's resume against the specified job role and requirements.

**Crucial Instruction: YOU MUST FILL ALL FIELDS IN THE JSON SCHEMA BELOW.**
If a piece of information is genuinely not found in the resume, explicitly state "Not Found" for string fields, "0/X" for scores, or an empty array `[]` for lists, but **DO NOT leave any field missing or null**.
The 'suggested_salary_range' should contain any salary mentioned in the resume or a placeholder like "To be calculated by system".

**Candidate Resume to Analyze:**
{resume_text}

**Job Role:** {job_role}
**Desired Experience:** {experience_info}

{job_description_section}

**JSON Fields (REQUIRED - MUST BE PRESENT):**
- **full_name**: string (e.g., "John Doe" or "Not Found")
- **contact_number**: string (e.g., "+1-555-123-4567" or "Not Found")
- **overall_experience**: string (e.g., "5 years 3 months" or "Not Found")
- **current_company_name**: string (e.g., "Acme Corp" or "Not Found")
- **current_company_address**: string (e.g., "New York, NY" or "Not Found")
- **hiring_recommendation**: string ("Hire", "Marginally Fit", or "Reject")
- **suggested_salary_range**: string (The salary found in the resume, or "To be calculated by system")
- **experience_match**: string ("Good Match", "Underqualified", "Overqualified", or "No Resume Provided")
- **analysis_summary**: object containing:
    - **candidate_overview**: string (detailed summary, use `\\n` for newlines)
    - **technical_prowess**: object (e.g., {{"Languages": ["Python"], "Tools": ["Git"]}} or {{}} if none)
    - **project_impact**: array of objects (each with "title", "company", "role", "achievements": ["string"] or [])
    - **education_certifications**: object (e.g., {{"education": ["B.Sc. Computer Science"], "certifications": ["AWS Certified"]}} or {{"education": [], "certifications": []}})
    - **overall_rating**: string (e.g., "4.5/5" or "Not Rated")
    - **conclusion**: string (overall concluding remarks, use `\\n` for newlines)
- **candidate_fitment_analysis**: object containing:
    - **strategic_alignment**: string
    - **comparable_experience_analysis**: string
    - **quantifiable_impact**: string
    - **potential_gaps_risks**: string
- **scoring_matrix**: array of 5 objects (each with "dimension", "weight", "score", "justification")
    - Ensure scores are in "x/Y" format.
- **aggregate_score**: string (e.g., "90/100" or "0/100")
- **fitment_verdict**: string ("SELECTED", "MARGINALLY FIT", or "REJECTED")
- **bench_recommendation**: object containing:
    - **retain**: string ("Yes", "No", or "Maybe")
    - **ideal_future_roles**: array of strings (e.g., ["Senior Data Scientist"] or [])
    - **justification**: string
- **alternative_role_recommendations**: array of objects (each with "role", "explanation")
- **automated_recruiter_insights**: object containing:
    - **red_flag_detection**: string (e.g., "Frequent job hopping" or "None detected")
    - **growth_indicators**: string (e.g., "Consistent upskilling" or "Not evident")
    - **cross_industry_potential**: string (e.g., "Transferable skills to finance" or "Limited")
- **interview_questions**: array of 7 relevant interview questions (strings).
"""
    
    input_variables = ["resume_text", "job_role", "experience_info", "job_description_section"]

    summary_prompt = PromptTemplate(
        input_variables=input_variables,
        template=summary_template
    )

    chain = LLMChain(llm=llm, prompt=summary_prompt)

    try:
        logging.info("Invoking LLMChain for resume analysis.")
        result = chain.invoke(input={
            "resume_text": resume_text,
            "job_role": job_role,
            "experience_info": experience_info,
            "job_description_section": job_description_section
        })
        llm_output = result["text"].strip()
        
        # ... (LLM output trimming and cleanup logic remains the same) ...
        json_match = re.search(r'\{.*\}', llm_output, re.DOTALL)
        if json_match:
            llm_output = json_match.group(0)
            llm_output = llm_output.replace('```json', '').replace('```', '').strip()
            llm_output = llm_output.replace('’', "'").replace('‘', "'").replace('“', '"').replace('”', '"')
        
        if not llm_output:
            logging.warning("LLM returned an empty response after cleanup.")
            return {"error": "LLM returned empty response. Cannot parse."}

        try:
            summary = json.loads(llm_output)
            logging.info("Successfully parsed LLM output as JSON.")
        except json.JSONDecodeError as e:
            # ... (Fallback JSON parsing logic remains the same) ...
            try:
                import ast
                summary = ast.literal_eval(llm_output)
                if not isinstance(summary, dict):
                    raise ValueError("`ast.literal_eval` did not result in a dictionary.")
                logging.warning("Used ast.literal_eval for JSON parsing due to initial failure.")
            except (ValueError, SyntaxError) as e_ast:
                logging.error(f"Error decoding LLM output even with fallback: {e_ast}")
                return {"error": f"Failed to decode LLM output: {e_ast}"}
        
        
        # ==========================================================
        # 🛑 CRITICAL FIX: POST-PROCESSING SALARY CALCULATION
        # ==========================================================
        
        # 1. Get experience from the LLM result
        overall_experience_str = summary.get("overall_experience", "Not Found")
        total_years = _parse_experience(overall_experience_str)
        
        # 2. Calculate the corrected salary range
        corrected_salary = _calculate_salary_range(job_role, total_years)
        
        # 3. OVERRIDE the LLM's suggested_salary_range
        summary["suggested_salary_range"] = corrected_salary
        logging.info(f"Overridden salary for {job_role} ({total_years:.1f} years): {corrected_salary}")
        
        # ==========================================================

        return summary
        
    except Exception as e:
        logging.error(f"An unexpected error occurred during LLM call: {e}")
        return {"error": f"An unexpected error occurred during AI analysis: {e}"}

# def analyze_resume_with_llm(resume_file_obj, job_description_file_obj, job_role, experience_type, min_years, max_years):
#     """
#     Orchestrates the entire resume analysis process.
#     1. Extracts text from the resume and job description files.
#     2. Builds a structured prompt for the LLM.
#     3. Calls the LLM (Langchain) to get the structured analysis.
#     4. Returns the parsed analysis data.
#     """
#     logging.info("Starting analyze_resume_with_llm function.")
    
#     # 1. Extract text from uploaded files
#     resume_text = extract_text_from_document(resume_file_obj, resume_file_obj.name)
    
#     job_description_text = ""
#     if job_description_file_obj:
#         job_description_text = extract_text_from_document(job_description_file_obj, job_description_file_obj.name)

#     if not resume_text:
#         logging.warning("No text extracted from resume. LLM analysis will be limited.")
#     else:
#         logging.info(f"Resume text extracted. Length: {len(resume_text)} chars.")
    
#     if not job_description_text:
#         logging.info("No text extracted from job description.")
#     else:
#         logging.info(f"JD text extracted. Length: {len(job_description_text)} chars.")

#     # Format experience info for the LLM prompt
#     experience_info = ""
#     if experience_type == "Specific Range (Years)":
#         experience_info = f"{min_years}-{max_years} years"
#     elif experience_type == "Minimum Years Required":
#         experience_info = f"at least {min_years} years"
#     else:
#         experience_info = experience_type

#     # 2. Call the LLM service function
#     analysis_data = llm_call(
#         resume_text=resume_text,
#         job_role=job_role,
#         experience_info=experience_info,
#         job_description_text=job_description_text
#     )
    
#     if analysis_data:
#         logging.info("LLM analysis data received successfully.")
#     else:
#         logging.error("LLM analysis data is None or empty after API call.")

#     return analysis_data


def analyze_resume_with_llm(resume_file_obj, job_description_file_obj=None, job_role="", experience_type="", min_years=0, max_years=0):
    """
    Handles both uploaded JD file and existing JD from CareerPage.
    """
    logging.info("Starting analyze_resume_with_llm function.")
    
    # 1. Extract text from resume
    resume_text = extract_text_from_document(resume_file_obj, resume_file_obj.name)
    
    # 2. Extract JD text
    job_description_text = ""
    if job_description_file_obj:
        if isinstance(job_description_file_obj, CareerPage):
            # Use the description field directly
            job_description_text = job_description_file_obj.description or ""
            logging.info(f"Using existing JD from CareerPage: {job_description_file_obj.title}")
        else:
            # Uploaded file
            job_description_text = extract_text_from_document(job_description_file_obj, job_description_file_obj.name)
            logging.info(f"JD text extracted from uploaded file. Length: {len(job_description_text)} chars.")
    
    # 3. Format experience info
    experience_info = ""
    if experience_type == "Specific Range (Years)":
        experience_info = f"{min_years}-{max_years} years"
    elif experience_type == "Minimum Years Required":
        experience_info = f"at least {min_years} years"
    else:
        experience_info = experience_type

    # 4. Call LLM
    analysis_data = llm_call(
        resume_text=resume_text,
        job_role=job_role,
        experience_info=experience_info,
        job_description_text=job_description_text
    )
    
    if analysis_data:
        logging.info("LLM analysis data received successfully.")
    else:
        logging.error("LLM analysis data is None or empty after API call.")

    return analysis_data


def make_blandai_call(phone_number, candidate_name, interview_questions_list):
    """
    Initiates an outbound call using Bland.ai for an interview.
    """
    if not BLANDAI_API_KEY:
        logging.error("BLANDAI_API_KEY is not set in environment variables.")
        return {"error": "BLANDAI_API_KEY is not set."}

    if interview_questions_list:
        script_core = "I have a few questions for you. " + " ".join(interview_questions_list[:3])
    else:
        script_core = "We'd like to discuss your application in more detail."

    headers = {
        "authorization": BLANDAI_API_KEY,
        "Content-Type": "application/json"
    }
    payload = {
        "phone_number": phone_number,
        "task": f"Hello {candidate_name}, this is an automated call from ATS regarding your application. Are you available for a brief automated interview? {script_core}",
        "voice_id": 9,
        "reduce_latency": True,
        "transfer_list": [],
        "answered_by_human": True,
        "first_sentence": f"Hello {candidate_name}, I am an AI interviewer calling to conduct a brief interview for the position you applied for. Are you ready to proceed?"
    }

    try:
        response = requests.post("https://api.bland.ai/v1/calls", headers=headers, json=payload)
        response.raise_for_status()
        logging.info(f"Bland.ai call initiated successfully. Response: {response.json()}")
        return response.json()
    except requests.exceptions.RequestException as e:
        logging.error(f"Bland.ai API call failed: {e}")
        return {"error": f"Bland.ai call failed: {e}"}

def get_blandai_call_details(call_id):
    """
    Fetches details of a Bland.ai call by its ID.
    """
    if not BLANDAI_API_KEY:
        logging.error("BLANDAI_API_KEY is not set in environment variables.")
        return {"error": "BLANDAI_API_KEY is not set."}

    headers = {
        "authorization": BLANDAI_API_KEY,
        "Content-Type": "application/json"
    }

    try:
        response = requests.get(f"https://api.bland.ai/v1/calls/{call_id}", headers=headers)
        response.raise_for_status()
        logging.info(f"Bland.ai call details for ID {call_id} fetched successfully.")
        return response.json()
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching Bland.ai call details for ID {call_id}: {e}")
        return {"error": f"Failed to fetch call details: {e}"}

def get_blandai_call_summary(call_id):
    """
    Fetches the summary of a Bland.ai call.
    """
    if not BLANDAI_API_KEY:
        logging.error("BLANDAI_API_KEY is not set in environment variables.")
        return {"error": "BLANDAI_API_KEY is not set."}

    headers = {
        "authorization": BLANDAI_API_KEY,
        "Content-Type": "application/json"
    }

    try:
        response = requests.get(f"https://api.bland.ai/v1/calls/{call_id}/summary", headers=headers)
        response.raise_for_status()
        logging.info(f"Bland.ai call summary for ID {call_id} fetched successfully.")
        return response.json()
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching Bland.ai call summary for ID {call_id}: {e}")
        return {"error": f"Failed to fetch call summary: {e}"}

##################### Complete Free ATS ####################


import re
from io import BytesIO
import logging
import spacy
from sentence_transformers import SentenceTransformer, util
from datetime import datetime
import PyPDF2 
import docx 
import os 

# --- Configuration & Setup (Unchanged) ---
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
SEMANTIC_MODEL_NAME = "all-mpnet-base-v2" 
SPACY_MODEL_NAME = "en_core_web_sm" 

nlp = None
embedding_model = None

try:
    nlp = spacy.load(SPACY_MODEL_NAME) 
    embedding_model = SentenceTransformer(SEMANTIC_MODEL_NAME)
    logging.info(f"Loaded Spacy model: {SPACY_MODEL_NAME} and SentenceTransformer model: {SEMANTIC_MODEL_NAME}")
except OSError as e:
    logging.error(f"Model loading failed. Ensure '{SPACY_MODEL_NAME}' is downloaded and files are accessible. Error: {e}")
    nlp = None
    embedding_model = None


# --- Constants (Unchanged) ---
SEMANTIC_WEIGHT = 0.1 
KEYWORD_WEIGHT = 0.9 
SCORE_SCALING_FACTOR = 3.55 
MATCH_THRESHOLD = 65 
WORK_EXPERIENCE_HEADINGS = [
    "work experience", "professional experience", "employment history", 
    "relevant experience", "experience", "job history"
]
EDUCATION_HEADINGS = [
    "education", "academic history", "educational qualifications", "qualifications"
]
# ------------------------------------------

# --- Existing Functions (Text Extraction, Candidate Info, Experience, Scoring) ---

# def extract_text(file_obj):
#     # ... (function body remains the same)
#     text = ""
#     fname = os.path.basename(getattr(file_obj, 'name', 'unknown_file')).lower()
    
#     try:
#         file_obj.seek(0)
        
#         if fname.endswith(".pdf"):
#             reader = PyPDF2.PdfReader(file_obj)
#             for page in reader.pages:
#                 page_text = page.extract_text()
#                 if page_text: text += page_text + "\n"
#         elif fname.endswith(".docx"):
#             doc = docx.Document(file_obj)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         else:
#             text = file_obj.read().decode('utf-8', errors='ignore')
            
#     except Exception as e:
#         logging.error(f"Text extraction failed for {fname}: {e}")
#         return ""
        
#     return re.sub(r'[^\S\r\n]+', ' ', text).strip()

def extract_text(file_obj):
    text = ""
    
    # Get filename if it exists
    fname = os.path.basename(getattr(file_obj, 'name', 'unknown_file')).lower()
    
    try:
        # Check if file_obj has 'seek' and 'read' methods (i.e., behaves like a file)
        if hasattr(file_obj, 'seek') and hasattr(file_obj, 'read'):
            file_obj.seek(0)
            
            if fname.endswith(".pdf"):
                reader = PyPDF2.PdfReader(file_obj)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            elif fname.endswith(".docx"):
                doc = docx.Document(file_obj)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:
                # For other file types, try to read as text
                raw_bytes = file_obj.read()
                # Decode bytes if needed
                if isinstance(raw_bytes, bytes):
                    text = raw_bytes.decode('utf-8', errors='ignore')
                else:
                    text = raw_bytes
        else:
            # If not a file-like object, maybe it's a model instance.
            # You can customize this block based on your model's fields
            
            # Example: If this is a CareerPage model, extract from text fields
            # Adjust these fields as per your actual model structure
            if hasattr(file_obj, 'title'):
                text += str(getattr(file_obj, 'title', '')) + "\n"
            if hasattr(file_obj, 'qualifications'):
                text += str(getattr(file_obj, 'qualifications', '')) + "\n"
            if hasattr(file_obj, 'responsibilities'):
                text += str(getattr(file_obj, 'responsibilities', '')) + "\n"
            if hasattr(file_obj, 'skills'):
                text += str(getattr(file_obj, 'skills', '')) + "\n"
            # Add any other fields you want to extract
            
            # If none of the above, fallback to empty string or str conversion
            if not text:
                text = str(file_obj)
    
    except Exception as e:
        logging.error(f"Text extraction failed for {fname}: {e}")
        return ""
    
    # Normalize whitespace (preserve new lines)
    return re.sub(r'[^\S\r\n]+', ' ', text).strip()



def is_header_or_title(text, job_role=""):
    # ... (function body remains the same)
    clean_text = text.strip().lower()
    
    common_headers_and_titles = {
        "resume", "curriculum vitae", "contact", "experience", "education",
        "profile", "summary", "objective", "skills", "declaration", "project",
        "career objective", "career summary", "references", "developer", 
        "engineer", "manager", "architect", "analyst", "specialist", "portfolio"
    }
    
    if job_role:
        common_headers_and_titles.add(job_role.lower())
        
    if len(clean_text.split()) > 5 or any(c.isdigit() or c in '()[]{}|@' for c in clean_text): 
        return True 
    
    if any(header in clean_text for header in common_headers_and_titles):
        if clean_text in common_headers_and_titles or job_role.lower() in clean_text:
             return True
            
    return False

import os, re


def extract_candidate_info(text, job_role="", resume_filename=""):
    """
    Extract full name, email, and phone number from resume text.
    Strong preference for real human-looking names.
    Fallback: derive clean name from email username.
    """

    # ------------------ EMAIL ------------------
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    email = email_match.group(0).strip() if email_match else "N/A"

    # ------------------ PHONE ------------------
    phone_match = re.search(
        r'(\+?\d{1,3}[\s\-\.]?)?\(?\d{2,4}\)?[\s\-\.]?\d{3,4}[\s\-\.]?\d{3,4}', text
    )
    phone = phone_match.group(0).strip() if phone_match else "N/A"

    # ------------------ TEXT CLEANUP ------------------
    lines = [re.sub(r'[^A-Za-z\s]', ' ', line).strip() for line in text.split("\n")]
    lines = [line for line in lines if line]

    # Words that *cannot* be a name
    banned_words = {
        "python", "django", "flask", "developer", "engineer", "java", "sql", "html",
        "css", "javascript", "react", "node", "git", "github", "aws", "docker",
        "api", "resume", "curriculum", "vitae", "profile", "objective", "skills",
        "education", "experience", "project"
    }

    probable_name = "N/A"

    # ------------------ STEP 1: spaCy PERSON ------------------
    if nlp is not None:
        doc = nlp(" ".join(lines[:50]))
        persons = [ent.text.strip() for ent in doc.ents if ent.label_ == "PERSON"]
        if persons:
            probable_name = persons[0]

    # ------------------ STEP 2: Manual header scan ------------------
    if probable_name == "N/A":
        for line in lines[:15]:
            if any(w.lower() in banned_words for w in line.lower().split()):
                continue
            words = line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
                probable_name = " ".join(words)
                break

    # ------------------ STEP 3: Validate name ------------------
    if (
        probable_name == "N/A"
        or len(probable_name) < 3
        or any(w.lower() in banned_words for w in probable_name.lower().split())
    ):
        probable_name = "N/A"

    # ------------------ STEP 4: Email-based fallback ------------------
    if probable_name == "N/A" and email != "N/A":
        username = email.split("@")[0]
        username = re.sub(r'[\d\._\-]+', ' ', username)
        username = re.sub(r'\s+', ' ', username).strip()

        parts = [p.capitalize() for p in username.split() if len(p) > 1]

        # if single chunk like "chiragmodi" → split mid into two
        if len(parts) == 1 and len(parts[0]) > 6:
            halves = re.findall(r'[A-Z]?[a-z]+', parts[0]) or [parts[0][:5], parts[0][5:]]
            probable_name = " ".join(h.capitalize() for h in halves[:2])
        elif len(parts) >= 2:
            probable_name = " ".join(parts[:2])
        elif len(parts) == 1:
            probable_name = parts[0]

    # ------------------ STEP 5: Filename fallback ------------------
    if probable_name == "N/A" and resume_filename:
        base = os.path.basename(resume_filename)
        base = re.sub(r'[_\-\.]', ' ', os.path.splitext(base)[0])
        probable_name = " ".join(w.capitalize() for w in base.split()[:2])

    # ------------------ FINAL CLEANUP ------------------
    probable_name = re.sub(r'\s+', ' ', probable_name).strip()
    if any(w.lower() in banned_words for w in probable_name.lower().split()):
        probable_name = "N/A"

    return probable_name or "N/A", email, phone


def extract_date_ranges(text):
    # ... (function body remains the same)
    date_patterns = [
        r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.,]?\s+\d{4})\s*[\-–]\s*(\b(?:Present|Current|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.,]?\s*\d{4}|\bPresent|\bCurrent)',
        r'(\d{4})\s*[\-–]\s*(\d{4}|\bPresent|\bCurrent)',
        r'(\b\d{1,2}/\d{4})\s*[\-–]\s*(\b\d{1,2}/\d{4}|\bPresent|\bCurrent)'
    ]
    
    date_ranges = []
    
    def parse_date(date_string):
        date_string = date_string.replace('Current', '').replace('Present', '').strip()
        if not date_string:
            return datetime.now()
        
        formats = ["%b %Y", "%B %Y", "%m/%Y", "%Y"]
        for fmt in formats:
            try:
                if len(date_string.split()) > 1 and date_string.split()[0].isalpha():
                    return datetime.strptime(date_string.split()[0][:3] + " " + date_string.split()[-1], "%b %Y")
                elif re.match(r'\d{4}', date_string) and len(date_string) == 4:
                    return datetime.strptime(date_string, "%Y")
                else:
                    return datetime.strptime(date_string, fmt)
            except ValueError:
                continue
        return None

    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for start_str, end_str in matches:
            start_date = parse_date(start_str)
            end_date = parse_date(end_str)
            
            if start_date and end_date and start_date < end_date:
                date_ranges.append((start_date, end_date))
                
    return date_ranges

WORK_EXPERIENCE_HEADINGS = [
    "work experience", "professional experience", "employment history",
    "relevant experience", "experience", "job history"
]

EDUCATION_HEADINGS = [
    "education", "academic history", "educational qualifications", "qualifications"
]

# --------------------- Helper functions ---------------------

def get_work_experience_text(resume_text):
    """
    Extract text between work experience heading and next section (education or end).
    """
    lines = resume_text.split("\n")
    experience_lines = []
    in_experience = False

    for line in lines:
        l = line.strip().lower()
        if any(h in l for h in WORK_EXPERIENCE_HEADINGS):
            in_experience = True
            continue
        if in_experience and any(h in l for h in EDUCATION_HEADINGS):
            break
        if in_experience:
            experience_lines.append(line)
    return "\n".join(experience_lines)


def parse_date_str(date_str):
    date_str = date_str.strip().replace(".", "")
    today = datetime.today()
    if re.search(r'present|current|now|till date', date_str, re.IGNORECASE):
        return today

    patterns = ["%b %Y", "%B %Y", "%m/%Y", "%Y"]
    for fmt in patterns:
        try:
            return datetime.strptime(date_str, fmt)
        except:
            continue
    return None





def calculate_duration_in_years_and_months(start, end):
    months = (end.year - start.year) * 12 + (end.month - start.month)
    years = months // 12
    rem_months = months % 12
    return years, rem_months


# --------------------- Main Function ---------------------

def extract_total_experience(resume_text):
    work_exp_text = get_work_experience_text(resume_text)
    if not work_exp_text.strip():
        return 0.0, "N/A"

    date_ranges = extract_date_ranges(work_exp_text)
    if not date_ranges:
        return 0.0, "N/A"

    # Sort & merge overlapping ranges
    date_ranges.sort(key=lambda x: x[0])
    merged = []
    cur_start, cur_end = date_ranges[0]

    for next_start, next_end in date_ranges[1:]:
        if next_start <= cur_end:
            cur_end = max(cur_end, next_end)
        else:
            merged.append((cur_start, cur_end))
            cur_start, cur_end = next_start, next_end
    merged.append((cur_start, cur_end))

    # Total duration
    total_months = 0
    for s, e in merged:
        y, m = calculate_duration_in_years_and_months(s, e)
        total_months += y * 12 + m

    total_years = total_months / 12
    y_int = int(total_years)
    m_int = int(round((total_years - y_int) * 12))
    experience_str = f"{y_int} Years, {m_int} Months"

    return round(total_years, 1), experience_str

def extract_required_experience(jd_text):
    # ... (function body remains the same)
    matches = re.findall(r'(\d+)(?:\s*[\+\-]\s*\d+)?\s*(?:to)?\s*years?', jd_text, re.IGNORECASE)
    
    if matches:
        min_exp = min(int(match[0]) for match in matches)
        return float(min_exp)
    
    min_match = re.search(r'(?:minimum|min|at least)\s*(\d+)\s*years?', jd_text, re.IGNORECASE)
    if min_match:
        return float(min_match.group(1))

    return 0.0 

def experience_match_score(candidate_years, required_years):
    # ... (function body remains the same)
    if required_years <= 0:
        return 100.0, "No Requirement Found"
        
    if candidate_years >= required_years:
        return 100.0, "Requirement Met or Exceeded"
    else:
        score = (candidate_years / required_years) * 100
        return round(min(score, 99.0), 1), f"Short by {round(required_years - candidate_years, 1)} years"

def semantic_similarity(resume_text, jd_text):
    # ... (function body remains the same)
    if nlp is None or embedding_model is None or not jd_text or not resume_text:
        return 0.0
    
    try:
        def filter_sentences(text):
            sents = [sent.text.lower() for sent in nlp(text).sents 
                     if len(sent.text.strip()) > 10 and any(c.isalpha() for c in sent.text)]
            return sents

        resume_sents = filter_sentences(resume_text)
        jd_sents = filter_sentences(jd_text)
        
        if not resume_sents or not jd_sents:
            logging.warning("Insufficient sentences for semantic analysis after filtering.")
            return 0.0

        emb_resume = embedding_model.encode(resume_sents, convert_to_tensor=True)
        emb_jd = embedding_model.encode(jd_sents, convert_to_tensor=True)
        
        cos_scores = util.cos_sim(emb_resume, emb_jd)
        
        max_scores_per_jd_sent = cos_scores.max(dim=0).values
        
        semantic_score = float(max_scores_per_jd_sent.mean() * 100)
        
        return round(semantic_score, 2)
        
    except Exception as e:
        logging.warning(f"Semantic similarity failed: {e}")
        return 0.0

def get_keywords(text):
    # ... (function body remains the same)
    if nlp is None: return set()

    doc = nlp(text.lower())
    keywords = set()
    
    irrelevant = {"etc", "description", "cv", "ltd", "solution", "university", "ability", "requirement", "pvt"}
    
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip()
        if 2 <= len(phrase.split()) <= 5 and len(phrase) > 3 and phrase not in irrelevant:
            keywords.add(phrase)
            
    for token in doc:
        if token.pos_ in ('NOUN', 'PROPN') and token.is_alpha and not token.is_stop and len(token.text) > 2:
            if token.text not in irrelevant:
                keywords.add(token.text)
            
    return keywords

def keyword_match_score(resume_text, jd_text):
    # ... (function body remains the same)
    jd_keywords = get_keywords(jd_text)
    resume_keywords = get_keywords(resume_text)
    
    if not jd_keywords:
        return 0.0

    matched_keywords = jd_keywords.intersection(resume_keywords)
    score = (len(matched_keywords) / len(jd_keywords)) * 100
    
    logging.info(f"JD Keywords extracted ({len(jd_keywords)}): {list(jd_keywords)[:5]}...")
    logging.info(f"Matched Keywords ({len(matched_keywords)}): {list(matched_keywords)[:5]}...")

    return round(score, 2)
    
# ----------------- NEW ALIGNMENT FUNCTION -----------------

def generate_strategic_alignment(aggregate_score, candidate_years, required_years, sem_score, keyword_score):
    """
    Creates the strategic, actionable summary based on all calculated metrics.
    """
    alignment = []

    # 1. Overall Recommendation
    if aggregate_score >= 85:
        alignment.append("✅ High Priority Candidate: The overall fitment score is excellent. Proceed directly to the interview stage.")
    elif aggregate_score >= MATCH_THRESHOLD: # 65
        alignment.append("🔶 Qualified Candidate: The score indicates a strong baseline match. Requires a screening interview to validate experience depth.")
    else:
        alignment.append("🛑 Low Priority Candidate: The score is below the threshold. Reroute only if the talent pipeline is scarce.")

    # 2. Experience Risk
    if required_years > 0 and candidate_years < required_years * 0.8:
        gap = round(required_years - candidate_years, 1)
        alignment.append(f"⚠️ Experience Gap: Candidate is short by {gap} years. Focus interview questions on transferable skills and project ownership to compensate.")
    elif candidate_years >= required_years:
        alignment.append("🟢 Experience Met: Candidate meets or exceeds the required professional experience.")
    else:
        alignment.append("🔵 Experience Context: No explicit experience requirement found in JD, or candidate is close to target.")
    
    # 3. Score Breakdown Insight
    if sem_score > keyword_score * 2:
        alignment.append("💡 Semantic Strength: The candidate's *concept* of the role aligns well (high semantic score). They may use different terminology than the JD. Review project descriptions closely.")
    elif keyword_score > sem_score * 1.5:
        alignment.append("💡 Keyword Focus: The resume is highly optimized for keywords. Ensure the candidate can articulate their experience (low semantic score suggests a potential knowledge gap beneath the surface.)")
    else:
        alignment.append("✅ Balanced Profile: Strong, consistent alignment in both skills (keywords) and conceptual understanding (semantic).")
        
    return "\n- " + "\n- ".join(alignment)


# ----------------- Overview Generation (Unchanged) -----------------
def generate_candidate_overview(name, job_role, aggregate_score, sem_score, keyword_score, recommendation, exp_match_score, req_exp):
    """Generates a descriptive, data-driven summary of the candidate's fit."""
    
    fit_level = ""
    if aggregate_score >= 90:
        fit_level = "an Exceptional Match"
        strength = "The alignment is near-perfect, indicating a strong cultural and technical fit."
    elif aggregate_score >= 80:
        fit_level = "a Strong Match"
        strength = "The candidate possesses a high degree of core competency and is immediately hirable."
    elif aggregate_score >= 65:
        fit_level = "a Moderate to Strong Match"
        strength = "The core skills are present, but some key areas may require further evaluation."
    else:
        fit_level = "a Borderline Match"
        strength = "The profile meets minimum conceptual requirements but lacks significant keyword depth."

    if sem_score > keyword_score * 2:
        match_type = "The high Semantic Score suggests strong conceptual understanding and relevant experience, even if specific keywords were not explicitly matched."
    elif keyword_score > sem_score * 1.5:
        match_type = "The profile is heavily keyword-optimized, but conceptual alignment is a concern. A manual review is essential to verify experience depth."
    else:
        match_type = "The conceptual and keyword matches are balanced, providing a consistent view of the candidate's profile."
        
    candidate_identifier = f"Mr./Ms. {name}" if name != "N/A" else "The Candidate"

    exp_context = ""
    if req_exp > 0 and exp_match_score < 100.0:
        exp_context = f"**Experience Alert:** Candidate only meets {exp_match_score:.1f}% of the required {req_exp} years of experience. "
    elif req_exp > 0:
        exp_context = f"Experience requirement of {req_exp} years is met. "


    overview = (
        f"{exp_context}"
        f"{candidate_identifier} is currently rated as {fit_level} for the {job_role} position "
        f"with an Aggregate Score of {aggregate_score:.2f}%. {strength} "
        f"The system recommends: {recommendation}. {match_type}"
    )
    return overview


# ----------------- ATS Analyzer (Main Function) -----------------
def analyze_resume(resume_file_obj, job_description_file_obj, job_role="general role"):
    if nlp is None or embedding_model is None:
        return {
            "full_name":"N/A",
            "contact_number":"N/A",
            "email":"N/A",
            "job_role": job_role,
            "aggregate_score":0.0,
            "fitment_verdict":"FATAL ERROR",
            "hiring_recommendation":"Not Applicable",
            "overall_experience": "N/A",
            "experience_match": "N/A",
            "strategic_alignment": "FATAL ERROR: Models failed to load.",
            "analysis_summary":"NLP/Embedding models failed to load. Cannot proceed.",
            "candidate_overview": "Model loading failed.",
            "overall_rating_summary": "N/A"
        }

    logging.info("Starting text extraction...")
    resume_text = extract_text(resume_file_obj)
    jd_text = extract_text(job_description_file_obj)
    
    if not resume_text or not jd_text: 
        return {
            "full_name":"N/A",
            "contact_number":"N/A",
            "email":"N/A",
            "job_role": job_role,
            "aggregate_score":0.0,
            "fitment_verdict":"ERROR",
            "hiring_recommendation":"Not Applicable",
            "overall_experience": "N/A",
            "experience_match": "N/A",
            "strategic_alignment": "ERROR: Text extraction failed.",
            "analysis_summary":"Failed to extract text from one or both files.",
            "candidate_overview": "Failed to extract necessary text for analysis.",
            "overall_rating_summary": "N/A"
        }

    # 1. Extract Candidate Info
    name, email, phone = extract_candidate_info(resume_text, job_role) 
    
    # 2. Extract and Score Experience
    candidate_years, candidate_exp_str = extract_total_experience(resume_text)
    required_years = extract_required_experience(jd_text)
    
    exp_match_score_val, exp_match_verdict = experience_match_score(candidate_years, required_years)
    
    overall_experience_display = candidate_exp_str
    experience_match_display = f"{exp_match_score_val}% ({exp_match_verdict}) [Req: {required_years} Yrs]"
    
    # 3. Calculate Scores (Semantic and Keyword)
    sem_score = semantic_similarity(resume_text, jd_text)
    keyword_score = keyword_match_score(resume_text, jd_text)

    # 4. Aggregate Score
    weighted_score = (sem_score * SEMANTIC_WEIGHT) + (keyword_score * KEYWORD_WEIGHT)
    aggregate_score = round(min(weighted_score * SCORE_SCALING_FACTOR, 100.0), 2)
    
    # 5. Calculate Overall Rating (0-5 scale)
    overall_rating_summary = f"{round((aggregate_score / 100) * 5, 1)}/5"

    logging.info(f"Final Scaled Score: {aggregate_score}% | Overall Rating: {overall_rating_summary}")

    # 6. Determine Verdict and Recommendation
    if aggregate_score >= 85: 
        verdict = "Selected / High Match"
        recommendation = "Hire Recommended"
    elif aggregate_score >= MATCH_THRESHOLD: # 65
        verdict = "Borderline / Review"
        recommendation = "Manual Review Recommended"
    else:
        verdict = "Fail / Not Selected"
        recommendation = "Not Recommended"

    # 7. Generate Strategic Alignment
    strategic_alignment = generate_strategic_alignment(
        aggregate_score, candidate_years, required_years, sem_score, keyword_score
    )

    # 8. Generate Descriptive Overview
    overview = generate_candidate_overview(
        name, job_role, aggregate_score, sem_score, keyword_score, 
        recommendation, exp_match_score_val, required_years
    )

    return {
        "full_name": name,
        "contact_number": phone,
        "email": email,
        "job_role": job_role,
        "overall_experience": overall_experience_display, 
        "experience_match": experience_match_display,     
        "aggregate_score": aggregate_score,
        "fitment_verdict": verdict,
        "hiring_recommendation": recommendation,
        "analysis_summary": (
            f"Semantic Score: {sem_score}%; Keyword Score: {keyword_score}%; "
            f"Final Scaled Score: {aggregate_score}%"
        ),
        "semantic_score": f"{sem_score}%",
        "keyword_score": f"{keyword_score}%",
        "final_scaled_score": f"{aggregate_score}%",
        "strategic_alignment": strategic_alignment,
        "candidate_overview": overview,
        "overall_rating_summary": overall_rating_summary  # ✅ Added
    }
