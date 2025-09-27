# Django core imports
from email.message import EmailMessage
import time
import PyPDF2
from django.http import HttpResponse, JsonResponse, HttpResponseBadRequest
from django.views.decorators.csrf import csrf_protect, csrf_exempt
from django.views.decorators.http import require_POST
from django.core.files.storage import default_storage, FileSystemStorage
from django.core.files.base import ContentFile
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.conf import settings
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required, user_passes_test
from django.template.defaulttags import register
from django.db.models import Q, Case, When, IntegerField
import docx
import google.generativeai as genai
# Standard libraries
import os
import re
import json
import random
import logging
import pytz
from datetime import datetime, timedelta
from collections import defaultdict
from urllib.parse import urljoin



# Third-party libraries
import requests
# import spacy
import docx2txt           # For Word documents
import fitz               # For PDFs (PyMuPDF)
import phonenumbers       # For phone number parsing
import google             # Required for genai
import google.api_core.exceptions
import os
import logging
from hr_app.UserCreationForm import AdminUserCreationForm
from hr_app.linkedin_multi_poster import post_jobs_to_linkedin
from hr_app.models import EmailConfiguration
import email
import imaplib
from hr_app.admin import User
navigation_logger = logging.getLogger('hr_app_navigation') #
# Local imports
from .forms import ResumeUploadForm, FinalDecisionForm, PhoneNumberForm, CustomUserCreationForm, CustomAuthenticationForm
from .models import Application, Apply_career, CandidateAnalysis, CareerAdvanceAnalysis, CareerJob, CareerPage, Category, CompanyInfo, Document, DraftEmail, Folder, JobApplicationFormSettings, JobDescriptionDocument, SentEmail, ThemeSettings
from .services import llm_call
from hr_app import services
from django.core.files.uploadedfile import SimpleUploadedFile
# import win32com.client
# import pythoncom
import string
from django.contrib.auth.hashers import make_password
from django.shortcuts import render, redirect
from django.urls import reverse
from django.http import HttpResponse
from django.core.mail import send_mail
from django.contrib import messages
# from .models import JobDescription # Assuming you have this model defined in models.py


from django.shortcuts import render, redirect, get_object_or_404, HttpResponse
from django.contrib import messages
from django.core.mail import EmailMessage
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
import json
import uuid
from django.contrib.auth.decorators import user_passes_test
from django.contrib.auth import get_user_model


from django.shortcuts import get_object_or_404
# from sentence_transformers import SentenceTransformer, util
import os

from django.utils.dateparse import parse_date
User = get_user_model() # Get the currently active user model
# Assuming these are already defined correctly
resume_storage = FileSystemStorage(location='media/resumes')
job_description_storage = FileSystemStorage(location='media/job_descriptions')

def signup_view(request):
    """
    Handles user registration using a Django form.
    """
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            # Save the new user and set a default role
            user = form.save(commit=False)
            user.role = 'user'  # Automatically set the role to 'user'
            user.save()
            
            # Log the user in after successful registration
            login(request, user)
            messages.success(request, "Account created successfully! Welcome.")
            return redirect('dashboard') # Redirect to your dashboard
        else:
            # Display form errors to the user
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"Error in {field}: {error}")
    else:
        form = CustomUserCreationForm()
        
    return render(request, 'signup.html', {'form': form})


def signin_view(request):
    """
    Handles user login using a Django authentication form for better validation.
    """
    if request.method == 'POST':
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            # The form's is_valid() method authenticates the user
            user = form.get_user()
            login(request, user)
            messages.success(request, f"Welcome back, {user.username}!")
            return redirect('dashboard')
        else:
            # The form handles the error message for invalid credentials
            messages.error(request, "Invalid username or password.")
    else:
        form = CustomAuthenticationForm()
    
    return render(request, 'signin.html', {'form': form})


@login_required
def signout_view(request):
    """
    Handles user logout.
    """
    logout(request)
    messages.info(request, "You have been logged out.")
    return redirect('signin') # Redirect to signin page after logout


# --- Role-based Access Control Helpers ---

def is_admin(user):
    """
    Checks if the user has 'admin' or 'superadmin' role.
    """
    return user.is_authenticated and (user.role == 'admin' or user.role == 'superadmin')

def is_superadmin(user):
    """
    Checks if the user has 'superadmin' role.
    """
    return user.is_authenticated and user.role == 'superadmin'


# --- Dashboard and existing views with login_required and role checks ---

# fs = FileSystemStorage(location='media/resumes')  # Save to 'media/resumes' folder
fs = FileSystemStorage(location='media/resumes', base_url='/media/resumes/')


resume_storage = FileSystemStorage(location=settings.RESUME_UPLOAD_ROOT, base_url=settings.RESUME_UPLOAD_URL)
job_description_storage = FileSystemStorage(location=settings.JOB_DESCRIPTION_UPLOAD_ROOT, base_url=settings.JOB_DESCRIPTION_UPLOAD_URL)


# @login_required # Ensure user is logged in to access this page
def home(request):
    return render(request, 'home.html')

def _parse_experience_string(experience_str):
    """
    Helper function to parse an experience string (e.g., "3 years", "5+", "10")
    into a numerical integer for comparison.
    Returns an integer or None if parsing fails.
    """
    if not isinstance(experience_str, str):
        return None # Not a string, cannot parse

    # Remove "years" or "year" and any leading/trailing whitespace
    clean_str = experience_str.lower().replace('years', '').replace('year', '').strip()

    # Handle "X+" cases (e.g., "5+", "10+")
    match_plus = re.match(r'(\d+)\+', clean_str)
    if match_plus:
        try:
            return int(match_plus.group(1)) # Return the number before the plus
        except ValueError:
            return None

    # Handle pure digit cases (e.g., "3", "5")
    if clean_str.isdigit():
        try:
            return int(clean_str)
        except ValueError:
            return None
            
    return None # Could not parse


def resume_analysis_view(request):
    analysis_result = None
    resume_url = None
    
    # Get all existing job description documents
    job_description_documents = JobDescriptionDocument.objects.all()

    form = ResumeUploadForm(request.POST or None, request.FILES or None)

    if request.method == 'POST':
        logging.info("POST request received for resume analysis.")
        
        # Check if an existing JD was selected from the dropdown
        existing_jd_id = request.POST.get('job_description_id')
        
        if form.is_valid():
            logging.info("Form is valid. Processing uploaded files and form data.")
            resume_file = form.cleaned_data['resume_pdf']
            
            # Use uploaded job description file first, if it exists
            if 'job_description' in request.FILES:
                job_description_file = form.cleaned_data['job_description']
            elif existing_jd_id:
                # If an existing JD was selected, fetch it from the database
                try:
                    job_description_doc = JobDescriptionDocument.objects.get(pk=existing_jd_id)
                    job_description_file = job_description_doc.file
                    logging.info(f"Using existing job description from database: {job_description_doc.title}")
                except JobDescriptionDocument.DoesNotExist:
                    messages.error(request, "Selected job description not found.")
                    logging.error(f"Job Description with ID {existing_jd_id} not found.")
                    job_description_file = None
            else:
                messages.error(request, "Please upload or select a job description.")
                job_description_file = None
            
            job_role = form.cleaned_data['job_role']
            target_experience_type = form.cleaned_data['target_experience_type']
            min_years_required = form.cleaned_data['min_years_required']
            max_years_required = form.cleaned_data['max_years_required']

            if job_description_file:
                try:
                    # 1. Save the resume file for PDF preview
                    resume_filename = resume_storage.save(resume_file.name, resume_file)
                    resume_url = request.build_absolute_uri(resume_storage.url(resume_filename))
                    logging.info(f"Resume file '{resume_filename}' saved for preview. URL: {resume_url}")

                    # 2. Call the main AI analysis service function.
                    llm_response = services.analyze_resume_with_llm(
                        resume_file_obj=resume_file,
                        job_description_file_obj=job_description_file,
                        job_role=job_role,
                        experience_type=target_experience_type,
                        min_years=min_years_required,
                        max_years=max_years_required
                    )

                    if llm_response and not llm_response.get("error"):
                        analysis_result = llm_response
                        
                        # --- START: DATABASE SAVE LOGIC ---
                        try:
                            # Get the analysis_summary dictionary safely
                            analysis_summary = analysis_result.get("analysis_summary", {})
                            candidate_fitment_analysis = analysis_result.get("candidate_fitment_analysis", {})
                            
                            # Prepare data for the CandidateAnalysis model.
                            # Serialize complex data to JSON strings.
                            candidate_data_for_db = {
                                "resume_file_path": resume_filename,
                                "full_name": analysis_result.get("full_name"),
                                "job_role": job_role,
                                "phone_no": analysis_result.get("contact_number"),
                                "hiring_recommendation": analysis_result.get("hiring_recommendation"),
                                "suggested_salary_range": analysis_result.get("suggested_salary_range"),
                                "interview_questions": json.dumps(analysis_result.get("interview_questions", [])),
                                "analysis_summary": json.dumps(analysis_summary),
                                "experience_match": analysis_result.get("experience_match"),
                                "overall_experience": analysis_result.get("overall_experience"),
                                "current_company_name": analysis_result.get("current_company_name"),
                                "current_company_address": analysis_result.get("current_company_address"),
                                "fitment_verdict": analysis_result.get("fitment_verdict"),
                                "aggregate_score": analysis_result.get("aggregate_score"),
                                "strategic_alignment": candidate_fitment_analysis.get("strategic_alignment", ""),
                                "quantifiable_impact": candidate_fitment_analysis.get("quantifiable_impact", ""),
                                "potential_gaps_risks": candidate_fitment_analysis.get("potential_gaps_risks", ""),
                                "comparable_experience": candidate_fitment_analysis.get("comparable_experience_analysis", ""),
                                "scoring_matrix_json": json.dumps(analysis_result.get("scoring_matrix", [])),
                                "bench_recommendation_json": json.dumps(analysis_result.get("bench_recommendation", {})),
                                "alternative_role_recommendations_json": json.dumps(analysis_result.get("alternative_role_recommendations", [])),
                                "automated_recruiter_insights_json": json.dumps(analysis_result.get("automated_recruiter_insights", {})),
                                "candidate_overview": analysis_summary.get("candidate_overview", ""),
                                "technical_prowess_json": json.dumps(analysis_summary.get("technical_prowess", {})),
                                "project_impact_json": json.dumps(analysis_summary.get("project_impact", [])),
                                "education_certifications_json": json.dumps(analysis_summary.get("education_certifications", {})),
                                "overall_rating_summary": analysis_summary.get("overall_rating", ""),
                                "conclusion_summary": analysis_summary.get("conclusion", ""),
                                # --- THE CRUCIAL LINE IS ADDED HERE ---
                                "user": request.user 
                            }
                            
                            # Use .create() to save the new object and get its automatically generated ID.
                            candidate_obj = CandidateAnalysis.objects.create(**candidate_data_for_db)
                            
                            # Now, update the analysis_result dictionary with the new ID
                            analysis_result['id'] = candidate_obj.id 
                            
                            messages.success(request, f"Analysis saved to database for {candidate_obj.full_name}.")
                        except Exception as db_save_error:
                            logging.warning(f"AI analysis completed, but failed to save to database: {db_save_error}")
                            messages.warning(request, f"AI analysis completed, but failed to save to database: {db_save_error}")
                        # --- END: DATABASE SAVE LOGIC ---
                        
                        messages.success(request, f"AI analysis completed for {analysis_result.get('full_name', 'the candidate')}.")
                    else:
                        error_message = llm_response.get("error", "AI analysis failed to return a valid response.") if llm_response else "LLM response was empty or None."
                        logging.error(f"LLM response error: {error_message}")
                        messages.error(request, error_message)
                except Exception as e:
                    logging.error(f"An unexpected error occurred during the analysis process: {e}", exc_info=True)
                    messages.error(request, f"An unexpected error occurred during analysis: {e}")
            
        else:
            logging.warning("Form is not valid. Displaying errors.")
            messages.error(request, "Please correct the errors in the form before submitting.")

    context = {
        'form': form,
        'analysis_result': analysis_result,
        'resume_url': resume_url,
        'job_description_documents': job_description_documents,
    }

    return render(request, 'resume_analysis.html', context)

@login_required
def interview_dashboard_view(request):
    """
    Displays a dashboard of all candidates with their high-level interview status
    for the currently logged-in user.
    """
    # Filter the initial querysets to only show data for the logged-in user
    user_candidates_query = CandidateAnalysis.objects.filter(user=request.user)
    
    unique_job_roles = user_candidates_query.values_list('job_role', flat=True).distinct().exclude(job_role__isnull=True).exclude(job_role__exact='').order_by('job_role')

    all_candidates_query = user_candidates_query.filter(interview_status='Pending')
    completed_interviews = user_candidates_query.filter(interview_status='Complete')

    selected_job_role = request.GET.get('job_role')
    if selected_job_role:
        all_candidates_query = all_candidates_query.filter(job_role=selected_job_role)

    try:
        all_candidates = all_candidates_query.order_by('-created_at')
    except Exception:  # Catch any potential FieldError
        all_candidates = all_candidates_query.order_by('-id')
        messages.warning(request, "Could not sort by 'created_at'. Sorting by creation order (ID) instead. Consider adding a 'created_at' or 'last_updated' field to your CandidateAnalysis model.")

    for candidate in all_candidates:
        if candidate.bland_call_id:
            try:
                # Make sure the 'services' module is imported correctly
                call_details = services.get_blandai_call_details(candidate.bland_call_id)
                if call_details and not call_details.get('error'):
                    candidate.call_details = call_details
                else:
                    candidate.call_details = {'status': 'error', 'error': call_details.get('error', 'API Error')}
            except Exception as e:
                candidate.call_details = {'status': 'error', 'error': f'Fetch failed: {e}'}
        else:
            candidate.call_details = None

    context = {
        'all_candidates': all_candidates,
        'unique_job_roles': unique_job_roles,
        'completed_interviews': completed_interviews,
    }
    return render(request, 'interview_dashboard.html', context)


@login_required
def candidate_profile(request):

    return render(request, 'candidate_profile.html') #

@login_required
def interview_detail_view(request, candidate_id):
    """
    Displays detailed interview status for a single candidate and allows triggering Bland.ai calls.
    This replaces much of the old 'interview_status_view' functionality and is linked from the dashboard tiles.
    """
    candidate = get_object_or_404(CandidateAnalysis, id=candidate_id) #
    phone_form = PhoneNumberForm(initial={'phone_number': candidate.phone_no}) #

    call_details = None #
    call_summary = None #

    if candidate.bland_call_id: #
        call_details = services.get_blandai_call_details(candidate.bland_call_id) #
        if call_details and not call_details.get('error'): #
            if call_details.get('status') == 'completed': #
                call_summary = services.get_blandai_call_summary(candidate.bland_call_id) #
                if call_summary and not call_summary.get('error'): #
                    messages.info(request, "Call completed and summary fetched.") #
                else:
                    messages.warning(request, call_summary.get('error', "Could not fetch call summary.")) #
            else:
                messages.info(request, f"Call status: {call_details.get('status', 'Unknown').replace('-', ' ').title()}. Refresh to check for updates.") #
        else:
            messages.warning(request, call_details.get('error', "Could not fetch call details for the existing call ID.")) #
    
    if request.method == 'POST': #
        action = request.POST.get('action') #

        if action == 'start_call': #
            phone_form = PhoneNumberForm(request.POST) #
            if phone_form.is_valid(): #
                phone_number = phone_form.cleaned_data['phone_number'] #
                
                if not phone_number and candidate.phone_no: #
                    phone_number = candidate.phone_no #
                
                if not phone_number: #
                    messages.error(request, "A valid phone number is required to start the call. It must start with '+' and country code, e.g., +919876543210.") #
                else:
                    if candidate.interview_questions: #
                        try: #
                            call_response = services.make_blandai_call(phone_number, candidate.full_name, json.loads(candidate.interview_questions)) #
                            if call_response and not call_response.get('error'): #
                                candidate.bland_call_id = call_response.get('call_id') #
                                candidate.save() #
                                messages.success(request, f"Call initiated successfully! Call ID: {candidate.bland_call_id}. Status: {call_response.get('status', 'Initiated').replace('-', ' ').title()}") #
                                return redirect('interview_detail', candidate_id=candidate.id) #
                            else:
                                messages.error(request, call_response.get('error', "Failed to initiate call with Bland.ai.")) #
                        except json.JSONDecodeError:
                            messages.error(request, "Error parsing interview questions. Please re-analyze the resume.")
                        except Exception as e: #
                            messages.error(request, f"An error occurred while initiating call: {e}") #
                    else:
                        messages.error(request, "No interview questions generated. Please analyze a resume first.") #
            else:
                messages.error(request, "Invalid phone number provided in the form.") #
        
        elif action == 'fetch_status' and candidate and candidate.bland_call_id: #
            return redirect('interview_detail', candidate_id=candidate.id) #

    context = { #
        'candidate': candidate, #
        'call_details': call_details, #
        'call_summary': call_summary, #
        'phone_form': phone_form, #
    }
    return render(request, 'interview_detail.html', context) #


@login_required
def interview_status_view(request, candidate_id=None):
    if candidate_id:
        single_candidate = get_object_or_404(CandidateAnalysis, id=candidate_id)
        single_call_details = None
        single_call_summary = None
        
        phone_form = PhoneNumberForm(initial={'phone_number': single_candidate.phone_no})
        
        # Manually create the FinalDecisionForm and populate it
        initial_decision = single_candidate.final_decision if single_candidate.final_decision else None
        initial_salary = single_candidate.final_salary if single_candidate.final_salary else None
        
        final_decision_form = FinalDecisionForm(initial={
            'final_decision': initial_decision,
            'final_salary': initial_salary
        })

        if request.method == 'POST':
            action = request.POST.get('action')

            if action == 'start_call':
                phone_form = PhoneNumberForm(request.POST)
                if phone_form.is_valid():
                    # ... your existing logic for starting a call ...
                    pass

            elif action == 'save_decision':
                final_decision_form = FinalDecisionForm(request.POST)
                if final_decision_form.is_valid():
                    # Manually save the data to the model instance
                    single_candidate.final_decision = final_decision_form.cleaned_data['final_decision']
                    single_candidate.final_salary = final_decision_form.cleaned_data['final_salary']
                    single_candidate.save()
                    messages.success(request, "Final decision and salary updated successfully!")
                else:
                    messages.error(request, "Error updating final decision. Please check your inputs.")
            
            return redirect('interview_status', candidate_id=single_candidate.id)

        if single_candidate.bland_call_id:
            single_call_details = services.get_blandai_call_details(single_candidate.bland_call_id)
            if single_call_details and single_call_details.get('status') == 'completed':
                single_call_summary = services.get_blandai_call_summary(single_candidate.bland_call_id)

        context = {
            'candidate': single_candidate,
            'call_details': single_call_details,
            'call_summary': single_call_summary,
            'phone_form': phone_form,
            'form': final_decision_form, # Pass your form instance to the template
        }
        return render(request, 'interview_status.html', context)
    else:
        # ... (your existing dashboard logic) ...
        all_candidates = CandidateAnalysis.objects.all()
        unique_job_roles = CandidateAnalysis.objects.values_list('job_role', flat=True).distinct().order_by('job_role')
        experience_options = [('1', '1 Year'), ('2', '2 Years'), ('3', '3 Years'), ('4', '4 Years'),
                              ('5', '5+ Years'), ('10+', '10+ Years')]

        job_role_filter = request.GET.get('job_role')
        experience_filter = request.GET.get('experience')

        if job_role_filter:
            all_candidates = all_candidates.filter(job_role=job_role_filter)
        if experience_filter:
            if experience_filter.isdigit():
                all_candidates = all_candidates.filter(overall_experience=int(experience_filter))
            elif experience_filter == '5+':
                all_candidates = all_candidates.filter(overall_experience__gte=5)
            elif experience_filter == '10+':
                all_candidates = all_candidates.filter(overall_experience__gte=10)
        
        context = {
            'all_candidates': all_candidates,
            'unique_job_roles': unique_job_roles,
            'experience_options': experience_options,
            'selected_job_role': job_role_filter,
            'selected_experience': experience_filter,
        }
        return render(request, 'interview_status.html', context)
    

def candidate_profile_view(request, candidate_id):
    """
    Displays a comprehensive profile for a single candidate.
    Fetches all related data including interview details and summary if available.
    Also handles POST requests for finalizing decision and salary, and interview status.
    """
    candidate = get_object_or_404(CandidateAnalysis, pk=candidate_id)
    
    call_details = None
    call_summary = None
    resume_url = None # Initialize resume_url

    # Generate resume_url if resume_file_path exists
    if hasattr(candidate, 'resume_file_path') and candidate.resume_file_path:
        try:
            resume_url = request.build_absolute_uri(resume_storage.url(candidate.resume_file_path))
        except Exception as e:
            logging.error(f"Error generating resume URL for candidate {candidate_id}: {e}")
            messages.error(request, "Could not generate resume preview URL.")


    # Handle POST request for forms on this page
    if request.method == 'POST':
        form_type = request.POST.get('form_type')
        print(f"DEBUG: POST request received. form_type: {form_type}") # Debugging line

        if form_type == 'finalize_decision_form':
            logging.info(f"POST request for finalize_decision_form received for candidate ID: {candidate_id}")
            
            final_decision = request.POST.get('final_decision')
            final_salary_str = request.POST.get('final_salary')
            interview_status = request.POST.get('interview_status') # NEW: Get interview status

            # Update candidate object fields
            candidate.final_decision = final_decision
            candidate.interview_status = interview_status # NEW: Update interview status
            
            # Convert final_salary to integer, handle potential errors
            if final_salary_str:
                try:
                    candidate.final_salary = int(final_salary_str)
                except ValueError:
                    messages.error(request, "Invalid salary amount. Please enter a valid number.")
                    logging.error(f"Invalid salary amount received for candidate {candidate_id}: {final_salary_str}")
                else: # Only save if conversion was successful
                    try:
                        candidate.save()
                        messages.success(request, f"Final decision, salary, and interview status saved successfully for {candidate.full_name}.")
                        logging.info(f"Final decision, salary, and interview status saved for candidate {candidate.pk}.")
                    except Exception as e:
                        messages.error(request, f"Error saving final decision, salary, and interview status: {e}")
                        logging.error(f"Error saving final decision, salary, and interview status for candidate {candidate.pk}: {e}", exc_info=True)
            else:
                candidate.final_salary = None # Set to None if empty
                try:
                    candidate.save() # Save even if salary is None
                    messages.success(request, f"Final decision and interview status saved successfully for {candidate.full_name}.")
                    logging.info(f"Final decision and interview status saved (salary cleared) for candidate {candidate.pk}.")
                except Exception as e:
                    messages.error(request, f"Error saving final decision and interview status: {e}")
                    logging.error(f"Error saving final decision and interview status for candidate {candidate.pk}: {e}", exc_info=True)
        
        elif form_type == 'initiate_interview_form':
            logging.info(f"POST request for initiate_interview_form received for candidate ID: {candidate_id}")
            phone_number = request.POST.get('phone_number')
            
            messages.info(request, f"Initiate AI Interview form submitted for {phone_number}. (Logic to be implemented)")

        else:
            messages.warning(request, "Unknown form submitted.")
            logging.warning(f"Unknown form_type received: {form_type}. Request POST data: {request.POST}")


    # Fetch Bland AI call details if a call ID exists (this part remains the same)
    if hasattr(candidate, 'bland_call_id') and candidate.bland_call_id:
        try:
            call_details = services.get_blandai_call_details(candidate.bland_call_id)
            if call_details and call_details.get('status') == 'completed':
                call_summary = services.get_blandai_call_summary(candidate.bland_call_id)
        except Exception as e:
            logging.error(f"Error fetching BlandAI call details or summary: {e}")
            messages.error(request, f"Error fetching AI interview details: {e}")

    # Initialize lists/dicts for JSON fields to ensure they are always available in context
    interview_questions_list = []
    scoring_matrix_list = []
    bench_recommendation_dict = {}
    alternative_role_recommendations_list = []
    automated_recruiter_insights_dict = {}
    technical_prowess_dict = {}
    project_impact_list = []
    education_certifications_dict = {}
    analysis_summary_dict = {} # Keep for compatibility with core_competencies_assessment

    # Attempt to parse JSON fields from the candidate object
    if candidate.interview_questions:
        try:
            interview_questions_list = json.loads(candidate.interview_questions)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing interview questions for this candidate.")
    
    if candidate.scoring_matrix_json:
        try:
            scoring_matrix_list = json.loads(candidate.scoring_matrix_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing scoring matrix for this candidate.")

    if candidate.bench_recommendation_json:
        try:
            bench_recommendation_dict = json.loads(candidate.bench_recommendation_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing bench recommendation for this candidate.")

    if candidate.alternative_role_recommendations_json:
        try:
            alternative_role_recommendations_list = json.loads(candidate.alternative_role_recommendations_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing alternative role recommendations for this candidate.")

    if candidate.automated_recruiter_insights_json:
        try:
            automated_recruiter_insights_dict = json.loads(candidate.automated_recruiter_insights_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing automated recruiter insights for this candidate.")

    if candidate.technical_prowess_json:
        try:
            technical_prowess_dict = json.loads(candidate.technical_prowess_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing technical prowess for this candidate.")

    if candidate.project_impact_json:
        try:
            project_impact_list = json.loads(candidate.project_impact_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing project impact for this candidate.")

    if candidate.education_certifications_json:
        try:
            education_certifications_dict = json.loads(candidate.education_certifications_json)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing education and certifications for this candidate.")

    # This is kept as analysis_summary_dict.core_competencies_assessment is still used in HTML
    if candidate.analysis_summary: 
        try:
            analysis_summary_dict = json.loads(candidate.analysis_summary)
        except json.JSONDecodeError:
            messages.error(request, "Error parsing analysis summary for this candidate.")


    context = {
        'candidate': candidate,
        'call_details': call_details,
        'call_summary': call_summary,
        'interview_questions_list': interview_questions_list,
        'analysis_summary_dict': analysis_summary_dict, # Keep for compatibility with core_competencies_assessment
        'scoring_matrix_list': scoring_matrix_list,
        'bench_recommendation_dict': bench_recommendation_dict,
        'alternative_role_recommendations_list': alternative_role_recommendations_list,
        'automated_recruiter_insights_dict': automated_recruiter_insights_dict,
        'technical_prowess_dict': technical_prowess_dict,
        'project_impact_list': project_impact_list,
        'education_certifications_dict': education_certifications_dict,
        'resume_url': resume_url, # Pass resume_url to the template
    }
    return render(request, 'candidate_profile.html', context)
@login_required
def initiate_call_interview(request):
    """
    Displays interview status and allows triggering Bland.ai calls.
    """
    candidate = None #
    call_details = None #
    call_summary = None #
    phone_form = PhoneNumberForm() #

    candidate_id = request.session.get('current_candidate_id') #
    if candidate_id: #
        candidate = get_object_or_404(CandidateAnalysis, id=candidate_id) #
        
        if candidate.phone_no: #
            phone_form = PhoneNumberForm(initial={'phone_number': candidate.phone_no}) #

        if candidate.bland_call_id: #
            call_details = services.get_blandai_call_details(candidate.bland_call_id) #
            if call_details and not call_details.get('error'): #
                if call_details.get('status') == 'completed': #
                    call_summary = services.get_blandai_call_summary(candidate.bland_call_id) #
                    if call_summary and not call_summary.get('error'): #
                        messages.info(request, "Call completed and summary fetched.") #
                    else:
                        messages.warning(request, call_summary.get('error', "Could not fetch call summary.")) #
                else:
                    messages.info(request, f"Call status: {call_details.get('status').replace('-', ' ').title()}. Refresh to check for updates.") #
            else:
                messages.warning(request, call_details.get('error', "Could not fetch call details for the existing call ID.")) #
    
    if request.method == 'POST': #
        action = request.POST.get('action') #

        if action == 'start_call' and candidate: #
            phone_form = PhoneNumberForm(request.POST) #
            if phone_form.is_valid(): #
                phone_number = phone_form.cleaned_data['phone_number'] #
                
                if not phone_number and candidate.phone_no: #
                    phone_number = candidate.phone_no #
                
                if not phone_number or not phone_form.is_valid_phone_number(phone_number): #
                    messages.error(request, "A valid phone number is required to start the call. It must start with '+' and country code, e.g., +919876543210.") #
                else:
                    if candidate.interview_questions: #
                        try:
                            questions = json.loads(candidate.interview_questions)
                            call_response = services.make_blandai_call(phone_number, candidate.full_name, questions) #
                            if call_response and not call_response.get('error'): #
                                candidate.bland_call_id = call_response.get('call_id') #
                                candidate.save() #
                                messages.success(request, f"Call initiated successfully! Call ID: {candidate.bland_call_id}. Status: {call_response.get('status').replace('-', ' ').title()}") #
                                return redirect('interview_status') #
                            else:
                                messages.error(request, call_response.get('error', "Failed to initiate call with Bland.ai.")) #
                        except json.JSONDecodeError:
                            messages.error(request, "Error parsing interview questions. Please re-analyze the resume.")
                        except Exception as e: #
                            messages.error(request, f"An error occurred while initiating call: {e}") #
                    else:
                        messages.error(request, "No interview questions generated. Please analyze a resume first.") #
            else:
                messages.error(request, "Invalid phone number provided in the form.") #
        
        elif action == 'fetch_status' and candidate and candidate.bland_call_id: #
            return redirect('interview_status') #

    context = { #
        'candidate': candidate, #
        'call_details': call_details, #
        'call_summary': call_summary, #
        'phone_form': phone_form, #
    }
    return render(request, 'initiate_call_interview.html', context) #


@login_required
def top_recommendations_view(request):
    """
    Displays top recommended candidates based on job role and AI analysis
    for the currently logged-in user.
    """
    navigation_logger.info(
        f"User (ID: {request.user.id if request.user.is_authenticated else 'Anonymous'}) "
        f"navigated to Top Recommendations Page. Path: {request.path}"
    )

    # Filter the base queryset by the logged-in user
    user_candidates_query = CandidateAnalysis.objects.filter(user=request.user)

    unique_job_roles = user_candidates_query.values_list('job_role', flat=True).distinct().exclude(job_role__isnull=True).exclude(job_role__exact='').order_by('job_role')
    
    recommended_candidates_query = user_candidates_query

    selected_job_role = request.GET.get('job_role')
    if selected_job_role:
        recommended_candidates_query = recommended_candidates_query.filter(job_role=selected_job_role)
        messages.info(request, f"Showing recommendations for job role: {selected_job_role}")
        navigation_logger.info(f"Filtered recommendations by job role: {selected_job_role}")
    else:
        messages.info(request, "Showing top candidates across all job roles.")
        navigation_logger.info("Viewing all recommendations (no job role filter).")

    # Apply AI-driven recommendation logic:
    # Prioritize 'Hire' recommendations first, then 'Good Match' for experience.
    recommended_candidates = recommended_candidates_query.order_by(
        # Candidates recommended for hiring first
        Case(
            When(hiring_recommendation='Hire', then=0),
            When(hiring_recommendation='Resign', then=1),
            When(hiring_recommendation='Reject', then=2),
            default=3,
            output_field=IntegerField(),
        ),
        # Then by experience match
        Case(
            When(experience_match='Good Match', then=0),
            When(experience_match='Overqualified', then=1),
            When(experience_match='Underqualified', then=2),
            default=3,
            output_field=IntegerField(),
        ),
        '-created_at' # Finally, by most recent analysis
    )

    context = {
        'recommended_candidates': recommended_candidates,
        'unique_job_roles': unique_job_roles,
        'selected_job_role': selected_job_role,
        'page_heading': 'Top Recommended Candidates'
    }
    return render(request, 'top_recommendations.html', context)




def convert_to_lpa(salary_string):
    """
    Converts a salary string (e.g., '₹50,000 - ₹65,000 per month') to LPA.
    Handles 'monthly' and 'LPA' formats.
    """
    if not salary_string:
        return "N/A"

    salary_string = salary_string.lower().replace(',', '')
    numbers = re.findall(r'\d+', salary_string)
    
    if not numbers:
        return "N/A"
    
    if 'lpa' in salary_string or 'annum' in salary_string:
        # Already in LPA, just clean and format
        min_salary = int(numbers[0])
        max_salary = int(numbers[1]) if len(numbers) > 1 else min_salary
        return f"₹{min_salary / 100000:.2f} - ₹{max_salary / 100000:.2f} LPA"
    
    if 'month' in salary_string or 'monthly' in salary_string:
        # Convert monthly to LPA
        min_monthly = int(numbers[0])
        min_lpa = (min_monthly * 12) / 100000
        
        max_lpa = None
        if len(numbers) > 1:
            max_monthly = int(numbers[1])
            max_lpa = (max_monthly * 12) / 100000
        
        if max_lpa:
            return f"₹{min_lpa:.2f} - ₹{max_lpa:.2f} LPA"
        else:
            return f"₹{min_lpa:.2f} LPA"

    # Assume it's a direct yearly amount without 'lpa' or 'annum' specified
    min_salary = int(numbers[0])
    min_lpa = min_salary / 100000

    if len(numbers) > 1:
        max_salary = int(numbers[1])
        max_lpa = max_salary / 100000
        return f"₹{min_lpa:.2f} - ₹{max_lpa:.2f} LPA"
    else:
        return f"₹{min_lpa:.2f} LPA"


@login_required
def candidate_records_view(request):
    """
    Displays all stored candidate analysis records for the logged-in user.
    Allows for final decision and salary updates.
    """
    # Fetch all candidates from the database for the current user
    candidates = CandidateAnalysis.objects.filter(user=request.user).order_by('-created_at')
    
    # Process each candidate record to convert the salary to LPA
    processed_candidates = []
    for candidate in candidates:
        candidate_data = {
            'id': candidate.pk, # Changed .id to .pk
            'full_name': candidate.full_name,
            'job_role': candidate.job_role,
            'overall_experience': candidate.overall_experience,
            'hiring_recommendation': candidate.hiring_recommendation,
            'suggested_salary_range': convert_to_lpa(candidate.suggested_salary_range), # Salary is converted here
            'final_decision': candidate.final_decision,
            'final_salary': candidate.final_salary,
            'analysis_type': candidate.analysis_type,
            'resume_file_path': candidate.resume_file_path,
        }
        processed_candidates.append(candidate_data)

    # Get unique job roles for the current user
    job_roles = candidates.values_list('job_role', flat=True).distinct().order_by('job_role')

    final_decision_forms = {}
    for candidate in candidates:
        initial_data = {
            'final_decision': candidate.final_decision,
            'final_salary': candidate.final_salary
        }
        final_decision_forms[candidate.pk] = FinalDecisionForm(prefix=f'decision_{candidate.pk}', initial=initial_data) # Changed .id to .pk

    if request.method == 'POST':
        candidate_id = request.POST.get('candidate_id')
        if candidate_id:
            # Important: Filter the get_object_or_404 call by user as well for security
            candidate = get_object_or_404(CandidateAnalysis, pk=candidate_id, user=request.user) # Changed id to pk
            form = FinalDecisionForm(request.POST, prefix=f'decision_{candidate.pk}') # Changed .id to .pk
            
            if form.is_valid():
                candidate.final_decision = form.cleaned_data['final_decision']
                candidate.final_salary = form.cleaned_data['final_salary']
                candidate.save()
                messages.success(request, f"Final decision updated for {candidate.full_name}.")
                return redirect('records')
            else:
                final_decision_forms[candidate.pk] = form # Changed .id to .pk
                messages.error(request, "Error updating final decision.")
        else:
            messages.error(request, "Invalid request to update record.")

    context = {
        'candidates': processed_candidates,  # Pass the new processed list
        'job_roles': job_roles,
        'final_decision_forms': final_decision_forms,
    }
    return render(request, 'records.html', context)

def selected_candidate(request):
    """
    Displays all stored candidate analysis records.
    Allows for final decision and salary updates.
    """
    # candidates = CandidateAnalysis.objects.all().order_by('-created_at') #
    candidates = CandidateAnalysis.objects.filter(final_decision='selected').order_by('-created_at')

    final_decision_forms = {} #

    for candidate in candidates: #
        initial_data = { #
            'final_decision': candidate.final_decision, #
            'final_salary': candidate.final_salary #
        }
        final_decision_forms[candidate.id] = FinalDecisionForm(prefix=f'decision_{candidate.id}', initial=initial_data) #

    if request.method == 'POST': #
        candidate_id = request.POST.get('candidate_id') #
        if candidate_id: #
            candidate = get_object_or_404(CandidateAnalysis, id=candidate_id) #
            form = FinalDecisionForm(request.POST, prefix=f'decision_{candidate.id}') #
            
            if form.is_valid(): #
                candidate.final_decision = form.cleaned_data['final_decision'] #
                candidate.final_salary = form.cleaned_data['final_salary'] #
                candidate.save() #
                messages.success(request, f"Final decision updated for {candidate.full_name}.") #
                return redirect('records') #
            else:
                final_decision_forms[candidate.id] = form #
                messages.error(request, "Error updating final decision.") #
        else:
            messages.error(request, "Invalid request to update record.") #

    context = { #
        'candidates': candidates, #
        'final_decision_forms': final_decision_forms, #
    }
    return render(request, 'selected_candidate.html', context) #

def rejected_candidate(request):
    """
    Displays all stored candidate analysis records.
    Allows for final decision and salary updates.
    """
    candidates = CandidateAnalysis.objects.filter(final_decision='Not Selected').order_by('-created_at')
    final_decision_forms = {} #

    for candidate in candidates: #
        initial_data = { #
            'final_decision': candidate.final_decision, #
            'final_salary': candidate.final_salary #
        }
        final_decision_forms[candidate.id] = FinalDecisionForm(prefix=f'decision_{candidate.id}', initial=initial_data) #

    if request.method == 'POST': #
        candidate_id = request.POST.get('candidate_id') #
        if candidate_id: #
            candidate = get_object_or_404(CandidateAnalysis, id=candidate_id) #
            form = FinalDecisionForm(request.POST, prefix=f'decision_{candidate.id}') #
            
            if form.is_valid(): #
                candidate.final_decision = form.cleaned_data['final_decision'] #
                candidate.final_salary = form.cleaned_data['final_salary'] #
                candidate.save() #
                messages.success(request, f"Final decision updated for {candidate.full_name}.") #
                return redirect('records') #
            else:
                final_decision_forms[candidate.id] = form #
                messages.error(request, "Error updating final decision.") #
        else:
            messages.error(request, "Invalid request to update record.") #

    context = { #
        'candidates': candidates, #
        'final_decision_forms': final_decision_forms, #
    }
    return render(request, 'rejected_candidate.html', context) #

def shortlisted_candidate(request):
    """
    Displays all stored candidate analysis records.
    Allows for final decision and salary updates.
    """
    candidates = CandidateAnalysis.objects.filter(final_decision='shortlisted').order_by('-created_at')
    final_decision_forms = {} #

    for candidate in candidates: #
        initial_data = { #
            'final_decision': candidate.final_decision, #
            'final_salary': candidate.final_salary #
        }
        final_decision_forms[candidate.id] = FinalDecisionForm(prefix=f'decision_{candidate.id}', initial=initial_data) #

    if request.method == 'POST': #
        candidate_id = request.POST.get('candidate_id') #
        if candidate_id: #
            candidate = get_object_or_404(CandidateAnalysis, id=candidate_id) #
            form = FinalDecisionForm(request.POST, prefix=f'decision_{candidate.id}') #
            
            if form.is_valid(): #
                candidate.final_decision = form.cleaned_data['final_decision'] #
                candidate.final_salary = form.cleaned_data['final_salary'] #
                candidate.save() #
                messages.success(request, f"Final decision updated for {candidate.full_name}.") #
                return redirect('records') #
            else:
                final_decision_forms[candidate.id] = form #
                messages.error(request, "Error updating final decision.") #
        else:
            messages.error(request, "Invalid request to update record.") #

    context = { #
        'candidates': candidates, #
        'final_decision_forms': final_decision_forms, #
    }
    return render(request, 'shortlisted_candidate.html', context) #

@login_required
def airtable_data_view(request):
    """
    Fetches and displays all data from the specified Airtable table.
    """
    records_df = services.fetch_airtable_data(table_name="Candidates") #
    
    if isinstance(records_df, dict) and records_df.get('error'): #
        messages.error(request, records_df.get('error')) #
        records_html = "<p>Error fetching Airtable data.</p>" #
    elif not records_df.empty: #
        records_html = records_df.to_html(classes="table table-striped table-bordered", index=False) #
    else:
        records_html = "<p>No records found in Airtable.</p>" #

    context = { #
        'records_html': records_html #
    }
    return render(request, 'hr_app/airtable_data.html', context) #

@login_required
def post_data_to_airtable_view(request):
    """
    Dynamically creates a form based on Airtable schema and allows posting data.
    """
    table_name = "Candidates" #
    schema_fields = services.get_airtable_schema(table_name) #
    
    if isinstance(schema_fields, dict) and schema_fields.get('error'): #
        messages.error(request, schema_fields.get('error')) #
        return render(request, 'hr_app/post_airtable.html', {'form': None, 'messages': messages.get_messages(request)}) #

    from django import forms as django_forms #
    
    DynamicAirtableForm = type('DynamicAirtableForm', (django_forms.Form,), {}) #

    for field in schema_fields: #
        field_name = field.get('name') #
        field_type = field.get('type') #
        if field_name: #
            if field_type == 'singleLineText' or field_type == 'multilineText': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.CharField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False), #
                    max_length=255 if field_type == 'singleLineText' else None, #
                    widget=django_forms.TextInput if field_type == 'singleLineText' else django_forms.Textarea #
                )
            elif field_type == 'number': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.DecimalField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False) #
                )
            elif field_type == 'checkbox': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.BooleanField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False), #
                    initial=False #
                )
            elif field_type == 'date': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.DateField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False), #
                    widget=django_forms.DateInput(attrs={'type': 'date'}) #
                )
            elif field_type == 'dateTime': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.DateTimeField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False), #
                    widget=django_forms.DateTimeInput(attrs={'type': 'datetime-local'}) #
                )
            elif field_type == 'email': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.EmailField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False) #
                )
            elif field_type == 'url': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.URLField( #
                    label=field_name.replace('_', ' ').title(), #
                    required=field.get('options', {}).get('isRequired', False) #
                )
            elif field_type == 'singleSelect': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.CharField( #
                    label=field_name.replace('_', ' ').title() + " (Single Select)", #
                    required=field.get('options', {}).get('isRequired', False), #
                    max_length=255 #
                )
            elif field_type == 'multipleSelect': #
                DynamicAirtableForm.base_fields[field_name] = django_forms.CharField( #
                    label=field_name.replace('_', ' ').title() + " (Comma-separated Multi Select)", #
                    required=field.get('options', {}).get('isRequired', False), #
                    help_text="Enter comma-separated values for multi-select fields." #
                )
            else:
                DynamicAirtableForm.base_fields[field_name] = django_forms.CharField( #
                    label=field_name.replace('_', ' ').title() + f" (Type: {field_type})", #
                    required=field.get('options', {}).get('isRequired', False), #
                    help_text=f"Field type '{field_type}' not explicitly handled. Enter as text." #
                )

    form = DynamicAirtableForm() #

    if request.method == 'POST': #
        form = DynamicAirtableForm(request.POST) #
        if form.is_valid(): #
            data_to_post = form.cleaned_data #
            
            for field in schema_fields: #
                if field.get('type') == 'multipleSelect' and field.get('name') in data_to_post: #
                    if isinstance(data_to_post[field.get('name')], str): #
                        data_to_post[field.get('name')] = [s.strip() for s in data_to_post[field.get('name')].split(',') if s.strip()] #
            
            final_data_for_airtable = {k: v for k, v in data_to_post.items() if v is not None and v != ''} #

            post_response = services.post_data_to_airtable(table_name, final_data_for_airtable) #
            if post_response and not post_response.get('error'): #
                messages.success(request, "Data successfully posted to Airtable!") #
                form = DynamicAirtableForm() #
            else:
                messages.error(request, post_response.get('error', "Failed to post data to Airtable.")) #
        else:
            messages.error(request, "Please correct the errors in the form.") #

    context = { #
        'form': form, #
        'table_name': table_name, #
    }
    return render(request, 'post_airtable.html', context) #
    
try: #
    import win32com.client #
    WIN32COM_AVAILABLE = True #
except ImportError: #
    WIN32COM_AVAILABLE = False #
    print("Warning: 'pywin32' library not found. Local Outlook access will not work.") #
    print("Please install it using: pip install pywin32") #


#['AIzaSyC2q_aKXeBkyZYUsHlX6_djPyyUTq126pc','AIzaSyBqmZT97k_5zoNyG2FB0xGG7nPoKf6fPkA']

GEMINI_API_KEY = "AIzaSyBqmZT97k_5zoNyG2FB0xGG7nPoKf6fPkA"
genai.configure(api_key=GEMINI_API_KEY)

model = genai.GenerativeModel('gemini-2.5-flash')



logger = logging.getLogger(__name__)



# try:
#     model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
# except Exception as e:
#     print(f"Error loading SentenceTransformer model: {e}")
#     model = None

# --- Existing Functions (Unchanged) ---
def calculate_basic_ats_score(resume_content, job_description_text):
    """
    Calculates the keyword-based match percentage.
    """
    if not job_description_text or not resume_content:
        return 0.0

    job_keywords = set(job_description_text.lower().split())
    resume_words = set(resume_content.lower().split())

    common_keywords = job_keywords.intersection(resume_words)
    keywords_found = [kw for kw in common_keywords if len(kw) > 3]

    match_percentage = (len(keywords_found) / len(job_keywords)) * 100 if job_keywords else 0
    return round(match_percentage, 2)

def extract_text_from_pdf(resume_path):
    try:
        doc = fitz.open(resume_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

def get_file_path(field_value):
    """
    Handles both FileField and string paths.
    Converts them into a usable absolute path under MEDIA_ROOT.
    """
    if not field_value:
        return None

    if hasattr(field_value, "name"):
        file_name = field_value.name
    else:
        file_name = str(field_value)

    file_name = file_name.replace("media/", "").lstrip("/")
    return os.path.normpath(os.path.join(settings.MEDIA_ROOT, file_name))

# --- MODIFIED: show_unread_emails ---
@login_required
@login_required
@login_required
def show_unread_emails(request):
    emails = Apply_career.objects.all().order_by('-date_applied')
    job_descriptions = JobDescriptionDocument.objects.all()

    processed_emails = []
    for email in emails:
        advanced_analysis_exists = CandidateAnalysis.objects.filter(
            application=email,
            analysis_type='Advance'
        ).exists()
        email.advanced_analysis_exists = advanced_analysis_exists

        matching_job_desc = job_descriptions.filter(title=email.career.title).first()
        if matching_job_desc:
            try:
                analysis, created = CandidateAnalysis.objects.get_or_create(
                    application=email,
                    analysis_type='Basic',
                    defaults={
                        'user': request.user,
                        'job_role': matching_job_desc.title,
                        'aggregate_score': 0.0,
                        'fitment_verdict': 'Pending'
                    }
                )

                if created:
                    resume_path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, email.resume.name.replace("media/", "").lstrip("/")))
                    resume_content = extract_text_from_pdf(resume_path)

                    job_desc_text = ""
                    if matching_job_desc.file:
                        job_desc_text = extract_text_from_pdf(matching_job_desc.file.path)
                    else:
                        job_desc_text = matching_job_desc.description

                    if model and resume_content and job_desc_text:
                        jd_embedding = model.encode(job_desc_text, convert_to_tensor=True)
                        resume_embedding = model.encode(resume_content, convert_to_tensor=True)
                        similarity_score = util.cos_sim(jd_embedding, resume_embedding).item() * 100
                        basic_score = round(similarity_score, 2)

                        analysis.aggregate_score = basic_score
                        analysis.save()
                    else:
                        basic_score = "N/A"
                else:
                    basic_score = analysis.aggregate_score

                email.selected_job_desc_id = matching_job_desc.id
                email.basic_ats_score = basic_score

            except Exception as e:
                print(f"Error during Basic ATS calculation: {e}")
                email.selected_job_desc_id = None
                email.basic_ats_score = "N/A"
        else:
            email.selected_job_desc_id = None
            email.basic_ats_score = "N/A"

        processed_emails.append(email)

    context = {
        'emails': processed_emails,
        'job_descriptions': job_descriptions,
    }

    return render(request, 'show_application.html', context)

# --- MODIFIED: basic_ats_analysis ---
def basic_ats_analysis(request, application_id, job_description_id):
    try:
        application = get_object_or_404(Apply_career, pk=application_id)
        job_description = get_object_or_404(JobDescriptionDocument, pk=job_description_id)
        
        # Check if an analysis record already exists for this application
        try:
            existing_analysis = CandidateAnalysis.objects.get(application=application)
            # If a record exists and has a score, return it immediately to avoid re-calculation
            if existing_analysis.aggregate_score is not None:
                match_percentage = existing_analysis.aggregate_score
                return JsonResponse({'success': True, 'analysis': {'match_percentage': match_percentage}})
        except CandidateAnalysis.DoesNotExist:
            # Continue if no record exists
            pass

        jd_text = ""
        jd_file_path = get_file_path(getattr(job_description, "file", None))
        jd_text = extract_text_from_file(jd_file_path)

        if getattr(job_description, "required_skills", None):
            jd_text += " " + job_description.required_skills
        if getattr(job_description, "preferred_skills", None):
            jd_text += " " + job_description.preferred_skills
        if getattr(job_description, "job_description", None):
            jd_text += " " + job_description.job_description

        if not jd_text.strip():
            return JsonResponse({'success': False, 'error': 'Job description text is empty.'}, status=400)

        resume_file_path = get_file_path(getattr(application, "resume", None))
        if not resume_file_path or not os.path.exists(resume_file_path):
            return JsonResponse({'success': False, 'error': 'Resume file not found.'}, status=400)

        resume_text = extract_text_from_pdf(resume_file_path)
        if not resume_text.strip():
            return JsonResponse({'success': False, 'error': 'Resume text is empty.'}, status=400)

        if model:
            jd_embedding = model.encode(jd_text, convert_to_tensor=True)
            resume_embedding = model.encode(resume_text, convert_to_tensor=True)
            similarity_score = util.cos_sim(jd_embedding, resume_embedding).item() * 100
            match_percentage = round(similarity_score, 2)
        else:
            match_percentage = 0.0

        # The key fix: Use update_or_create with a single lookup key (the OneToOneField)
        analysis_obj, created = CandidateAnalysis.objects.update_or_create(
            application=application,
            defaults={
                'user': request.user,
                'job_role': job_description.title,
                'aggregate_score': match_percentage,
                'fitment_verdict': 'Pending',
                'analysis_type': 'Basic',  # This field will be updated on every click
            }
        )
        
        analysis_data = {
            'match_percentage': match_percentage,
        }

        return JsonResponse({'success': True, 'analysis': analysis_data})

    except (Apply_career.DoesNotExist, JobDescriptionDocument.DoesNotExist):
        return JsonResponse({'success': False, 'error': 'Application or Job Description not found.'}, status=404)

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
# --- Existing Functions (Unchanged) ---
def advance_ats_analysis(request, application_id):
    """
    Performs advanced ATS analysis on a candidate's resume and saves the results to the database.
    This function now handles both POST (generate) and GET (view) requests.
    """
    try:
        application = get_object_or_404(Apply_career, pk=application_id)
        
        # Handle GET request to view existing data
        if request.method == 'GET':
            try:
                # Use .get() for OneToOneField relationships as there should be only one
                existing_analysis = CandidateAnalysis.objects.get(application=application)
                analysis_data = {
                    "full_name": existing_analysis.full_name,
                    "contact_number": existing_analysis.phone_no,
                    "overall_experience": existing_analysis.overall_experience,
                    "current_company_name": existing_analysis.current_company_name,
                    "current_company_address": existing_analysis.current_company_address,
                    "hiring_recommendation": existing_analysis.hiring_recommendation,
                    "suggested_salary_range": existing_analysis.suggested_salary_range,
                    "experience_match": existing_analysis.experience_match,
                    "analysis_summary": json.loads(existing_analysis.analysis_summary) if existing_analysis.analysis_summary else {},
                    "candidate_fitment_analysis": {
                        "strategic_alignment": existing_analysis.strategic_alignment,
                        "comparable_experience_analysis": existing_analysis.comparable_experience,
                        "quantifiable_impact": existing_analysis.quantifiable_impact,
                        "potential_gaps_risks": existing_analysis.potential_gaps_risks
                    },
                    "scoring_matrix": json.loads(existing_analysis.scoring_matrix_json) if existing_analysis.scoring_matrix_json else [],
                    "aggregate_score": existing_analysis.aggregate_score,
                    "fitment_verdict": existing_analysis.fitment_verdict,
                    "bench_recommendation": json.loads(existing_analysis.bench_recommendation_json) if existing_analysis.bench_recommendation_json else {},
                    "alternative_role_recommendations": json.loads(existing_analysis.alternative_role_recommendations_json) if existing_analysis.alternative_role_recommendations_json else {},
                    "automated_recruiter_insights": json.loads(existing_analysis.automated_recruiter_insights_json) if existing_analysis.automated_recruiter_insights_json else {},
                    "interview_questions": json.loads(existing_analysis.interview_questions) if existing_analysis.interview_questions else [],
                }
                return JsonResponse({'success': True, 'analysis': analysis_data})
            except CandidateAnalysis.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Analysis not found for this application.'}, status=404)
        
        # Handle POST request to generate new data
        elif request.method == 'POST':
            if not application.resume:
                return JsonResponse({'success': False, 'error': 'Resume not uploaded for this application.'}, status=400)
            
            resume_name = application.resume.name.replace("media/", "").lstrip("/")
            resume_path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, resume_name))
            
            if not os.path.exists(resume_path):
                return JsonResponse({'success': False, 'error': f'Resume file not found on disk at: {resume_path}'}, status=404)
            
            resume_content = extract_text_from_pdf(resume_path)
            if not resume_content:
                return JsonResponse({'success': False, 'error': 'Failed to extract text from resume.'}, status=400)
            
            job_description_section = ""
            job_title = "N/A"
            experience_info = "Not Specified"
            if application.career and application.career.description:
                job_description_section = f"""**Job Description to Analyze:**\n{application.career.description}"""
                job_title = application.career.title or "N/A"
            
            prompt = f"""
            **YOUR RESPONSE MUST BE A SINGLE, VALID JSON OBJECT. NO OTHER TEXT, NO MARKDOWN FENCES (```json), NO EXPLANATIONS.**
            You are an expert HR evaluator. Your primary goal is to provide a **comprehensive and complete JSON analysis** of the candidate's resume against the specified job role and requirements.
            **Crucial Instruction: YOU MUST FILL ALL FIELDS IN THE JSON SCHEMA BELOW.**
            If a piece of information is genuinely not found in the resume, explicitly state "Not Found" for string fields, "0/X" for scores, or an empty array `[]` for lists, but **DO NOT leave any field missing or null**.
            **Salary Guidelines:**
            - Less than 2 years experience: ₹15,000 to ₹25,000
            - 2 to 5 years experience: ₹35,000 to ₹65,000
            - More than 5 years experience: ₹50,000 to ₹80,000
            **Candidate Resume to Analyze:**
            {resume_content}
            **Job Role:** {job_title}
            **Desired Experience:** {experience_info}
            {job_description_section}
            **JSON Fields (REQUIRED - MUST BE PRESENT):**
            - **full_name**: string (e.g., "John Doe" or "Not Found")
            - **contact_number**: string (e.g., "+1-555-123-4567" or "Not Found")
            - **overall_experience**: string (e.g., "5 years 3 months" or "Not Found")
            - **current_company_name**: string (e.g., "Acme Corp" or "Not Found")
            - **current_company_address**: string (e.g., "New York, NY" or "Not Found")
            - **hiring_recommendation**: string ("Hire", "Marginally Fit", or "Reject")
            - **suggested_salary_range**: string (e.g., "₹8 LPA - ₹12 LPA" or "Not Applicable (No Experience)")
            - **experience_match**: string ("Good Match", "Underqualified", "Overqualified", or "No Resume Provided")
            - **analysis_summary**: object containing:
                - **candidate_overview**: string (detailed summary, use `\\n` for newlines)
                - **technical_prowess**: object (e.g., {{"Languages": ["Python"], "Tools": ["Git"]}} or {{}} if none)
                - **project_impact**: array of objects (each with "title", "company", "role", "achievements": ["string"] or [])
                - **education_certifications**: object (e.g., {{"education": ["B.Sc. Computer Science"], "certifications": ["AWS Certified"]}} or {{"education": [], "certifications": []}})
                - **overall_rating**: string (e.g., "4.5/5" or "Not Rated")
                - **conclusion**: string (overall concluding remarks, use `\\n` for newlines)
            - **candidate_fitment_analysis**: object containing:
                - **strategic_alignment**: string
                - **comparable_experience_analysis**: string
                - **quantifiable_impact**: string
                - **potential_gaps_risks**: string
            - **scoring_matrix**: array of 5 objects (each with "dimension", "weight", "score", "justification")
                - Ensure scores are in "x/Y" format.
            - **aggregate_score**: string (e.g., "90/100" or "0/100")
            - **fitment_verdict**: string ("SELECTED", "MARGINALLY FIT", or "REJECTED")
            - **bench_recommendation**: object containing:
                - **retain**: string ("Yes", "No", or "Maybe")
                - **ideal_future_roles**: array of strings (e.g., ["Senior Data Scientist"] or [])
                - **justification**: string
            - **alternative_role_recommendations**: array of objects (each with "role", "explanation")
            - **automated_recruiter_insights**: object containing:
                - **red_flag_detection**: string (e.g., "Frequent job hopping" or "None detected")
                - **growth_indicators**: string (e.g., "Consistent upskilling" or "Not evident")
                - **cross_industry_potential**: string (e.g., "Transferable skills to finance" or "Limited")
            - **interview_questions**: array of 7 relevant interview questions (strings).
            """
            
            # Make the dynamic call to the generative model.
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            
            try:
                clean_response = response.text.replace('```json', '').replace('```', '').strip()
                llm_analysis = json.loads(clean_response)
                
                analysis_summary = llm_analysis.get('analysis_summary', {})
                candidate_fitment = llm_analysis.get('candidate_fitment_analysis', {})
                
                # The OneToOneField is the lookup key. analysis_type is a detail field.
                analysis_obj, created = CandidateAnalysis.objects.update_or_create(
                    application=application,
                    defaults={
                        'user': request.user,
                        'job_role': job_title,
                        'resume_file_path': resume_path,
                        'analysis_type': 'Advance', # This is a detail, not part of the unique lookup.
                        'full_name': llm_analysis.get('full_name', 'Not Found'),
                        'phone_no': llm_analysis.get('contact_number', 'Not Found'),
                        'overall_experience': llm_analysis.get('overall_experience', 'Not Found'),
                        'current_company_name': llm_analysis.get('current_company_name', 'Not Found'),
                        'current_company_address': llm_analysis.get('current_company_address', 'Not Found'),
                        'hiring_recommendation': llm_analysis.get('hiring_recommendation', 'Not Found'),
                        'suggested_salary_range': llm_analysis.get('suggested_salary_range', 'Not Found'),
                        'experience_match': llm_analysis.get('experience_match', 'Not Found'),
                        'fitment_verdict': llm_analysis.get('fitment_verdict', 'Not Found'),
                        'aggregate_score': llm_analysis.get('aggregate_score', '0/100'),
                        'analysis_summary': json.dumps(analysis_summary),
                        'candidate_overview': analysis_summary.get('candidate_overview', 'Not Found'),
                        'technical_prowess_json': json.dumps(analysis_summary.get('technical_prowess', {})),
                        'project_impact_json': json.dumps(analysis_summary.get('project_impact', [])),
                        'education_certifications_json': json.dumps(analysis_summary.get('education_certifications', {})),
                        'overall_rating_summary': analysis_summary.get('overall_rating', 'Not Rated'),
                        'conclusion_summary': analysis_summary.get('conclusion', 'Not Found'),
                        'strategic_alignment': candidate_fitment.get('strategic_alignment', 'Not Found'),
                        'comparable_experience': candidate_fitment.get('comparable_experience_analysis', 'Not Found'),
                        'quantifiable_impact': candidate_fitment.get('quantifiable_impact', 'Not Found'),
                        'potential_gaps_risks': candidate_fitment.get('potential_gaps_risks', 'Not Found'),
                        'scoring_matrix_json': json.dumps(llm_analysis.get('scoring_matrix', [])),
                        'bench_recommendation_json': json.dumps(llm_analysis.get('bench_recommendation', {})),
                        'alternative_role_recommendations_json': json.dumps(llm_analysis.get('alternative_role_recommendations', [])),
                        'automated_recruiter_insights_json': json.dumps(llm_analysis.get('automated_recruiter_insights', {})),
                        'interview_questions': json.dumps(llm_analysis.get('interview_questions', []))
                    }
                )

                return JsonResponse({'success': True, 'analysis': llm_analysis})
            
            except (json.JSONDecodeError, ValueError) as e:
                print(f"Error parsing LLM response: {e}")
                return JsonResponse({'success': False, 'error': f'Failed to parse LLM analysis response. Error: {e}'}, status=500)
    
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return JsonResponse({'success': False, 'error': f'An unexpected error occurred: {str(e)}'}, status=500)
@require_POST
def update_application_data(request):
    try:
        data = json.loads(request.body)
        application_id = data.get('application_id')
        remark = data.get('remark')
        job_role = data.get('job_role')

        if not application_id:
            return JsonResponse({'success': False, 'message': 'Application ID is required.'}, status=400)
        
        try:
            application = Application.objects.get(id=application_id)
        except Application.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Application not found.'}, status=404)

        application.remark = remark
        application.job_role = job_role
        application.save()

        return JsonResponse({'success': True, 'message': 'Data saved successfully.'})

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data.'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': f'An error occurred: {str(e)}'}, status=500)
# In your views.py, update the analyze_resume function
def analyze_resume(request, email_id, analysis_type, job_description_id=None):
    """
    This function analyzes the resume content and returns a JSON response.
    This is where your logic for parsing and AI calls lives.
    """
    try:
        application = Application.objects.get(id=email_id)

        # Determine the resume file path
        if not application.resume_url:
            return JsonResponse({'error': 'Resume file not found for this application.'}, status=404)

        resume_file_path = os.path.join(settings.BASE_DIR, application.resume_url.lstrip('/'))

        # Read the resume content
        resume_content = ""
        if resume_file_path.endswith('.pdf'):
            try:
                doc = fitz.open(resume_file_path)
                resume_content = "".join(page.get_text() for page in doc)
                doc.close()
            except Exception as e:
                logger.error(f"Error reading PDF resume: {e}")
                return JsonResponse({'error': f"Error reading PDF resume: {e}"}, status=500)
        elif resume_file_path.endswith('.docx'):
            try:
                resume_content = docx2txt.process(resume_file_path)
            except Exception as e:
                logger.error(f"Error reading DOCX resume: {e}")
                return JsonResponse({'error': f"Error reading DOCX resume: {e}"}, status=500)

        # Get the Job Description content if an ID is provided
        job_description_content = ""
        job_description_doc = None
        if job_description_id:
            try:
                job_description_doc = JobDescriptionDocument.objects.get(id=job_description_id)
                jd_file_path = os.path.join(settings.MEDIA_ROOT, 'job_descriptions', os.path.basename(job_description_doc.file.name))

                if jd_file_path.endswith('.pdf'):
                    try:
                        doc = fitz.open(jd_file_path)
                        job_description_content = "".join(page.get_text() for page in doc)
                        doc.close()
                    except Exception as e:
                        logger.error(f"Error reading PDF job description: {e}")
                elif jd_file_path.endswith('.docx'):
                    try:
                        job_description_content = docx2txt.process(jd_file_path)
                    except Exception as e:
                        logger.error(f"Error reading DOCX job description: {e}")
            except JobDescriptionDocument.DoesNotExist:
                logger.warning(f"Job Description with ID {job_description_id} not found.")

        # Determine job role and experience information
        job_role = job_description_doc.title if job_description_id else 'General'
        experience_info = "Not specified" # This should be dynamically parsed from JD for a real app

        # Call the new service function for both analysis types
        analysis_data = llm_call(
            resume_text=resume_content,
            job_role=job_role,
            experience_info=experience_info,
            job_description_text=job_description_content
        )

        # Check if llm_call returned an error dictionary
        if "error" in analysis_data:
            return JsonResponse(analysis_data, status=500)

        # Create a new CandidateAnalysis object to store the results
        new_analysis = CandidateAnalysis.objects.create(
            full_name=application.candidate_name,
            job_role=job_role,
            resume_file_path=application.resume_url,
            analysis_summary=json.dumps(analysis_data)
        )

        # Redirect to the analysis results page
        return redirect('analysis_results', analysis_id=new_analysis.id)

    except Application.DoesNotExist:
        return JsonResponse({'error': 'Application not found.'}, status=404)
    except Exception as e:
        logger.error(f"Error analyzing resume: {e}")
        return JsonResponse({'error': f"An error occurred during analysis: {e}"}, status=500)

# def show_unread_emails(request):
#     return render(request, 'show_application.html')

@login_required
def process_ats_option(request, email_id, ats_type):
    """
    Handles the selection of Basic ATS or Premium ATS for a given email.
    This is a placeholder for actual ATS processing logic.
    """
    # In a real application, you would:
    # 1. Retrieve the email/application details based on email_id.
    # 2. Trigger the appropriate ATS analysis (Basic or Premium).
    # 3. Store the results in your database (e.g., CandidateAnalysis model).
    # 4. Redirect to a relevant page (e.g., candidate profile or an analysis report).

    if ats_type == 'basic':
        messages.success(request, f"Basic ATS processing initiated for email ID: {email_id}. (Placeholder)")
        logging.info(f"Basic ATS triggered for email ID: {email_id}")
    elif ats_type == 'premium':
        messages.info(request, f"Premium ATS processing initiated for email ID: {email_id}. (Placeholder)")
        logging.info(f"Premium ATS triggered for email ID: {email_id}")
    else:
        messages.error(request, f"Invalid ATS type: {ats_type}")
        logging.warning(f"Invalid ATS type received: {ats_type} for email ID: {email_id}")
    
    # Redirect back to the applications list or a confirmation page
    return redirect('show_applications')



@login_required
def dashboard(request):
    """
    Renders the HR Talent Scout Dashboard with various analytics and metrics
    derived from CandidateAnalysis data, dynamically fetching call details.
    """
    # 1. Fetch all candidates
    all_candidates = CandidateAnalysis.objects.all()

    # 2. Aggregate Data for Metrics and Charts by iterating through candidates
    # This is necessary because call_details is not a direct model field for filtering
    total_candidates = 0
    completed_calls = 0
    in_progress_calls = 0
    failed_calls = 0
    not_initiated_calls = 0

    job_role_counts = defaultdict(int)
    interview_status_counts = defaultdict(int)
    hiring_recommendation_counts = defaultdict(int)
    experience_buckets = defaultdict(int)

    for candidate in all_candidates:
        total_candidates += 1
        
        # Aggregate Job Role
        if candidate.job_role:
            job_role_counts[candidate.job_role] += 1

        # Aggregate Experience Level - NOW USING THE HELPER FUNCTION
        raw_exp_str = candidate.overall_experience
        exp = _parse_experience_string(raw_exp_str) # Parse the string to an integer
        
        if exp is not None: # Ensure parsed experience is not null
            if exp <= 1:
                experience_buckets['1 Year'] += 1
            elif exp == 2:
                experience_buckets['2 Years'] += 1
            elif exp == 3:
                experience_buckets['3 Years'] += 1
            elif exp == 4:
                experience_buckets['4 Years'] += 1
            elif 5 <= exp < 10: # Experience between 5 and 9 years
                experience_buckets['5-9 Years'] += 1
            else: # Experience 10 years or more (exp >= 10)
                experience_buckets['10+ Years'] += 1

        # Aggregate Interview Status and Call Metrics
        if candidate.bland_call_id:
            # Dynamically fetch call details for each candidate with a call ID
            # IMPORTANT: This might be slow if you have many candidates.
            # Consider caching or fetching in bulk if your Bland.ai API allows.
            call_details = services.get_blandai_call_details(candidate.bland_call_id)
            if call_details and not call_details.get('error'):
                status = call_details.get('status')
                if status == 'completed':
                    completed_calls += 1
                    interview_status_counts['completed'] += 1
                elif status == 'in_progress':
                    in_progress_calls += 1
                    interview_status_counts['in_progress'] += 1
                elif status == 'failed':
                    failed_calls += 1
                    interview_status_counts['failed'] += 1
                else:
                    # Catch any other statuses (e.g., 'queued', 'ringing')
                    interview_status_counts[status] += 1
            else:
                # If call_details fetch failed, treat as not initiated or failed to fetch
                interview_status_counts['failed_to_fetch'] += 1 # New category for dashboard
        else:
            not_initiated_calls += 1
            interview_status_counts['not_initiated'] += 1

        # Aggregate Hiring Recommendation
        if candidate.hiring_recommendation:
            hiring_recommendation_counts[candidate.hiring_recommendation] += 1

    # Calculate Call Success Rate
    total_calls_attempted = completed_calls + in_progress_calls + failed_calls
    success_rate = (completed_calls / total_calls_attempted * 100) if total_calls_attempted > 0 else 0
    success_rate = round(success_rate, 2)

    # 3. Prepare Chart Data Structures

    # Chart 1: Candidates by Job Role
    job_role_chart_data = {
        'labels': list(job_role_counts.keys()),
        'data': list(job_role_counts.values()),
    }

    # Chart 2: Interview Status Distribution
    interview_status_labels = []
    interview_status_data = []
    ordered_statuses = ['completed', 'in_progress', 'failed', 'not_initiated', 'failed_to_fetch'] # Include new status
    status_display_names = {
        'completed': 'Completed',
        'in_progress': 'In Progress',
        'failed': 'Failed',
        'not_initiated': 'Not Initiated',
        'failed_to_fetch': 'Call Data Error' # Display name for failed to fetch
    }
    for status_key in ordered_statuses:
        if interview_status_counts[status_key] > 0:
            interview_status_labels.append(status_display_names[status_key])
            interview_status_data.append(interview_status_counts[status_key])

    interview_status_chart_data = {
        'labels': interview_status_labels,
        'data': interview_status_data,
    }

    # Chart 3: Hiring Recommendation Distribution
    hiring_recommendation_chart_data = {
        'labels': list(hiring_recommendation_counts.keys()),
        'data': list(hiring_recommendation_counts.values()),
    }

    # Chart 4: Candidates by Experience Level
    ordered_experience_labels = ['1 Year', '2 Years', '3 Years', '4 Years', '5-9 Years', '10+ Years']
    experience_chart_labels = []
    experience_chart_data_values = []
    for label in ordered_experience_labels:
        if experience_buckets[label] > 0:
            experience_chart_labels.append(label)
            experience_chart_data_values.append(experience_buckets[label])

    experience_chart_data = {
        'labels': experience_chart_labels,
        'data': experience_chart_data_values,
    }

    # 4. Prepare Context
    context = {
        # Key Metrics
        'total_candidates': total_candidates,
        'completed_calls': completed_calls,
        'in_progress_calls': in_progress_calls,
        'failed_calls': failed_calls,
        'success_rate': success_rate,

        # Chart Data
        'job_role_chart_data': job_role_chart_data,
        'interview_status_chart_data': interview_status_chart_data,
        'hiring_recommendation_chart_data': hiring_recommendation_chart_data,
        'experience_chart_data': experience_chart_data,
    }

    return render(request, 'dashboard.html', context)


User = get_user_model() # Get the currently active user model

# Helper functions to check user roles
def is_superadmin(user):
    """Checks if the user has the 'superadmin' role."""
    return user.is_authenticated and user.role == 'superadmin'

def is_admin(user):
    """Checks if the user has the 'admin' role."""
    return user.is_authenticated and user.role == 'admin'

def is_admin_or_superadmin(user):
    """Checks if the user is either an 'admin' or 'superadmin'."""
    return is_admin(user) or is_superadmin(user)


@user_passes_test(is_superadmin, login_url='/signin/')
def superadmin_dashboard_view(request):
    # Get all users to display in the table
    all_users = User.objects.all()

    # Get the available roles from your model
    available_roles = User.ROLE_CHOICES # Assuming ROLE_CHOICES is defined in your model

    context = {
        'users': all_users,
        'available_roles': available_roles,
    }
    return render(request, 'superadmin_dashboard.html', context)


def admin_dashboard_view(request):
    """
    Renders the admin dashboard and handles user creation.
    """
    if request.method == 'POST':
        # --- Handle User Creation Form Submission ---
        username = request.POST.get('username')
        password = request.POST.get('password')
        role = request.POST.get('role', 'user')

        try:
            # Ensure only superadmins can create admin users
            if request.user.role == 'admin' and role != 'user':
                messages.error(request, "You do not have permission to create this type of user.")
            else:
                new_user = User.objects.create_user(username=username, password=password, role=role)
                if request.user.role == 'admin':
                    new_user.created_by = request.user
                
                # Ensure is_staff status for admin/superadmin roles
                if role == 'admin' or role == 'superadmin':
                    new_user.is_staff = True

                new_user.save()
                messages.success(request, f"User '{new_user.username}' created successfully with role '{new_user.get_role_display()}'.")
        except Exception as e:
            messages.error(request, f"Error creating user: {e}")

        # Always redirect after a POST request to prevent form resubmission
        return redirect('admin_dashboard') # Assuming your URL name is 'admin_dashboard'

    else: # request.method == 'GET'
        # --- Handle Page Load (Displaying Users) ---
        
        # Filter users based on the current user's role
        if request.user.role == 'superadmin':
            all_users = User.objects.all().order_by('-date_joined')
        else:
            all_users = User.objects.filter(created_by=request.user).order_by('-date_joined')

        context = {
            'users': all_users,
            'current_user_role': request.user.role,
            'current_username': request.user.username,
        }
        
        return render(request, 'admin_dashboard.html', context)

@login_required
def all_job_descriptions(request):
    """
    Renders the page to display all job postings created by the logged-in user.
    """
    # Filter the queryset to only show job postings belonging to the current user
    # Renamed variable to be more clear
    job_postings = CareerPage.objects.filter(user=request.user) 
    context = {
        'job_postings': job_postings,  # Updated context key
    }
    return render(request, 'all_job_descriptions.html', context)


@login_required
def create_job_description(request):
    """
    Handles the creation of a new job description via text input with detailed IT fields.
    The new job description will be linked to the current user.
    """
    if request.method == 'POST':
        title = request.POST.get('title')
        company_name = request.POST.get('company_name')
        job_level = request.POST.get('job_level')
        department = request.POST.get('department')
        country = request.POST.get('country')
        state = request.POST.get('state')
        city = request.POST.get('city')
        employment_type = request.POST.get('employment_type')
        salary_min = request.POST.get('salary_min')
        salary_max = request.POST.get('salary_max')
        salary_frequency = request.POST.get('salary_frequency')
        overview = request.POST.get('overview')
        responsibilities = request.POST.get('responsibilities')
        required_skills = request.POST.get('required_skills')
        preferred_skills = request.POST.get('preferred_skills')
        education_experience = request.POST.get('education_experience')
        benefits = request.POST.get('benefits')

        salary_min = int(salary_min) if salary_min else None
        salary_max = int(salary_max) if salary_max else None

        if title:
            description_text_content = f"Title: {title}\n"
            if company_name: description_text_content += f"Company: {company_name}\n"
            if job_level: description_text_content += f"Job Level: {job_level.replace('_', ' ').title()}\n"
            if department: description_text_content += f"Department: {department}\n"
            if country: description_text_content += f"Location: {country}"
            if state: description_text_content += f", {state}"
            if city: description_text_content += f", {city}\n"
            if employment_type: description_text_content += f"Employment Type: {employment_type.replace('-', ' ').title()}\n"
            if salary_min and salary_max: description_text_content += f"Salary: ${salary_min} - ${salary_max} ({salary_frequency})\n"
            elif salary_min: description_text_content += f"Salary: ${salary_min} ({salary_frequency})\n"
            elif salary_max: description_text_content += f"Salary: ${salary_max} ({salary_frequency})\n"
            if overview: description_text_content += f"\nOverview:\n{overview}\n"
            if responsibilities: description_text_content += f"\nResponsibilities:\n{responsibilities}\n"
            if required_skills: description_text_content += f"\nRequired Skills:\n{required_skills}\n"
            if preferred_skills: description_text_content += f"\nPreferred Skills:\n{preferred_skills}\n"
            if education_experience: description_text_content += f"\nEducation & Experience:\n{education_experience}\n"
            if benefits: description_text_content += f"\nBenefits:\n{benefits}\n"

            # Create a JobDescriptionDocument instance and link it to the current user
            JobDescriptionDocument.objects.create(
                user=request.user,  # This is the crucial line
                title=title,
                company_name=company_name,
                job_level=job_level,
                department=department,
                country=country,
                state=state,
                city=city,
                employment_type=employment_type,
                salary_min=salary_min,
                salary_max=salary_max,
                salary_frequency=salary_frequency,
                overview=overview,
                responsibilities=responsibilities,
                required_skills=required_skills,
                preferred_skills=preferred_skills,
                education_experience=education_experience,
                benefits=benefits,
                job_description=description_text_content,
            )
            messages.success(request, 'Job description created successfully!')
            return redirect('all_job_descriptions')
        else:
            messages.error(request, 'Please provide at least a title for the job description.')

    return render(request, 'create_job_description.html')







def upload_job_description(request):
    """
    Handles the uploading of a new job description file.
    """
    if request.method == 'POST':
        title = request.POST.get('title')
        uploaded_file = request.FILES.get('file')

        if title and uploaded_file:
            # Check if the file is a PDF or DOCX (or other allowed types)
            if uploaded_file.content_type in ['application/pdf', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'text/plain']:
                # Assign the user here!
                job_description = JobDescriptionDocument(
                    title=title,
                    file=uploaded_file,
                    user=request.user  # <--- Add this line
                )
                job_description.save()
                messages.success(request, 'Job description uploaded successfully!')
                return redirect('all_job_descriptions')
            else:
                messages.error(request, 'Invalid file type. Please upload a PDF, DOCX, or TXT file.')
        else:
            messages.error(request, 'Please provide both a title and a file.')

    return render(request, 'upload_job_description.html')

def edit_job_description(request, jd_id):
        """
        Handles editing an existing job description.
        You'll need to implement the form and saving logic here.
        """
        jd = get_object_or_404(JobDescriptionDocument, id=jd_id)
        if request.method == 'POST':
            # Handle form submission to update the JD
            jd.title = request.POST.get('title', jd.title)
            jd.job_level = request.POST.get('job_level', jd.job_level)
            jd.department = request.POST.get('department', jd.department)
            jd.location = request.POST.get('location', jd.location)
            jd.employment_type = request.POST.get('employment_type', jd.employment_type)
            jd.overview = request.POST.get('overview', jd.overview)
            jd.responsibilities = request.POST.get('responsibilities', jd.responsibilities)
            jd.required_skills = request.POST.get('required_skills', jd.required_skills)
            jd.preferred_skills = request.POST.get('preferred_skills', jd.preferred_skills)
            jd.education_experience = request.POST.get('education_experience', jd.education_experience)
            jd.benefits = request.POST.get('benefits', jd.benefits)
            # Handle file update if necessary (more complex, might involve deleting old and saving new)
            
            jd.save()
            messages.success(request, 'Job description updated successfully!')
            return redirect('analyze_jd', jd_id=jd.id) # Redirect to the detail page or all JDs

        context = {'jd': jd}
        return render(request, 'edit_job_description.html', context) # You'll create this template



@require_POST
def delete_jd(request, jd_id):
    """
    Deletes a job description document.
    """
    try:
        jd = get_object_or_404(JobDescriptionDocument, id=jd_id)
        # Delete the file from storage if it exists
        if jd.file:
            jd.file.delete(save=False) # save=False prevents an extra database save
        jd.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

def analyze_jd(request, jd_id):
    """
    Renders the page to display the full details of a job description.
    """
    jd = get_object_or_404(JobDescriptionDocument, id=jd_id)
    context = {'jd': jd}
    return render(request, 'analyze_jd.html', context)



def analyze_application_view(request, email_id, analysis_type, jd_id=None):
    """
    This function analyzes the resume content and returns a JSON response.
    This is where your logic for parsing and AI calls lives.
    """
    try:
        application = Application.objects.get(id=email_id)

        if not application.resume_url:
            return JsonResponse({'error': 'Resume file not found for this application.'}, status=404)

        resume_file_path = os.path.join(settings.BASE_DIR, application.resume_url.lstrip('/'))

        resume_content = ""
        if resume_file_path.endswith('.pdf'):
            try:
                doc = fitz.open(resume_file_path)
                resume_content = "".join(page.get_text() for page in doc)
                doc.close()
            except Exception as e:
                logger.error(f"Error reading PDF resume: {e}")
                return JsonResponse({'error': f"Error reading PDF resume: {e}"}, status=500)
        elif resume_file_path.endswith('.docx'):
            try:
                resume_content = docx2txt.process(resume_file_path)
            except Exception as e:
                logger.error(f"Error reading DOCX resume: {e}")
                return JsonResponse({'error': f"Error reading DOCX resume: {e}"}, status=500)

        job_description_content = ""
        job_title = "General"
        if jd_id:
            try:
                job_description_doc = JobDescriptionDocument.objects.get(id=jd_id)
                job_title = job_description_doc.job_title
                jd_file_path = os.path.join(settings.MEDIA_ROOT, 'job_descriptions', os.path.basename(job_description_doc.file.name))

                if jd_file_path.endswith('.pdf'):
                    try:
                        doc = fitz.open(jd_file_path)
                        job_description_content = "".join(page.get_text() for page in doc)
                        doc.close()
                    except Exception as e:
                        logger.error(f"Error reading PDF job description: {e}")
                elif jd_file_path.endswith('.docx'):
                    try:
                        job_description_content = docx2txt.process(jd_file_path)
                    except Exception as e:
                        logger.error(f"Error reading DOCX job description: {e}")
            except JobDescriptionDocument.DoesNotExist:
                logger.warning(f"Job Description with ID {jd_id} not found.")

        prompt = ""
        if analysis_type == 'basic':
            if job_description_content:
                prompt = f"""
                You are a resume screening bot. Please compare the following resume to the job description and provide a summary of the candidate's strengths and weaknesses for the role.
                Resume: {resume_content}
                Job Description: {job_description_content}
                """
            else:
                prompt = f"""
                You are an expert HR assistant. Analyze the following resume and provide a summary of the candidate's skills, experience, and key qualifications.
                Resume: {resume_content}
                """
        elif analysis_type == 'advanced':
            if not job_description_content:
                return JsonResponse({'error': 'Advanced analysis requires a job description.'}, status=400)
                
            prompt = f"""
            You are an advanced ATS (Applicant Tracking System) assistant. Perform a comprehensive analysis of the candidate's resume against the provided job description.
            
            Resume: {resume_content}
            
            Job Description: {job_description_content}
            
            Provide a detailed JSON object with the following keys:
            - 'match_score': An integer from 1 to 100 representing the relevance of the resume to the job description.
            - 'missing_keywords': A list of key skills or terms from the job description that were not found in the resume.
            - 'matching_keywords': A list of key skills or terms from the job description that were successfully found.
            - 'fit_summary': A detailed paragraph summarizing how well the candidate's experience and skills align with the job requirements.
            - 'areas_for_improvement': A list of suggestions for the candidate to improve their resume for this specific role.
            """

        # Perform the Gemini API call with the constructed prompt
        # Assuming `model` object from `genai.GenerativeModel` is available
        response = model.generate_content(prompt)
        
        if analysis_type == 'advanced':
            try:
                # Use a regular expression to find the JSON object and remove any surrounding text or formatting.
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    json_string = json_match.group(0)
                    analysis_data = json.loads(json_string)
                else:
                    raise json.JSONDecodeError("No JSON object found in response.", response.text, 0)
            except json.JSONDecodeError as e:
                # This is the improved error handling block
                logger.error(f"Failed to decode JSON from AI response: {response.text}. Error: {e}")
                return JsonResponse({
                    'error': 'Failed to parse AI response. The response was not valid JSON.',
                    'raw_ai_response': response.text
                }, status=500)
        else:
            analysis_data = {'summary': response.text}
            
        new_analysis = CandidateAnalysis.objects.create(
            full_name=application.candidate_name,
            job_role=job_title,
            resume_file_path=application.resume_url,
            analysis_summary=json.dumps(analysis_data)
        )

        return redirect('analysis_results', analysis_id=new_analysis.id)

    except Application.DoesNotExist:
        return JsonResponse({'error': 'Application not found.'}, status=404)
    except Exception as e:
        logger.error(f"An error occurred during analysis: {e}")
        return JsonResponse({'error': f"An error occurred during analysis: {e}"}, status=500)
        
def analysis_results_view(request, analysis_id):
    """
    Displays the results of a stored analysis.
    """
    analysis = get_object_or_404(CandidateAnalysis, pk=analysis_id)
    
    try:
        # Check for and remove the markdown code block formatting
        summary_text = analysis.analysis_summary
        if summary_text.startswith('```json\n') and summary_text.endswith('\n```'):
            json_string = summary_text.strip('`\n') # Remove backticks and newlines
            json_string = json_string.lstrip('json').strip() # Remove the 'json' label
            result_data = json.loads(json_string)
        else:
            # If no markdown, assume it's already a clean JSON string
            result_data = json.loads(summary_text)
            
    except (json.JSONDecodeError, TypeError) as e:
        logger.error(f"Error decoding JSON for analysis_id {analysis_id}: {e}")
        return HttpResponseBadRequest("Error: Analysis data is corrupted.")

    context = {
        'analysis': analysis,
        'result_data': result_data
    }
    return render(request, 'analysis_results.html', context)


###################
def basic_resume_analysis_view(request):
    """
    This Django view handles the AI-based resume analysis process.
    It can be triggered automatically from a URL or manually via a form.

    Args:
        request: The Django HttpRequest object.

    Returns:
        A Django HttpResponse that renders the analysis results page.
    """
    analysis_result = None
    resume_url = None
    job_description_documents = JobDescriptionDocument.objects.all()

    # Get application ID and job description ID from URL query parameters
    application_id = request.GET.get('application_id')
    job_description_id = request.GET.get('job_description_id')
    
    application = None
    if application_id:
        try:
            application = Application.objects.get(id=application_id)
            if application.resume_url:
                resume_url = application.resume_url
        except Application.DoesNotExist:
            messages.error(request, "Application not found.")
            logging.error(f"Application with ID {application_id} not found.")

    # Check if analysis should be performed automatically
    if application_id and job_description_id:
        logging.info("Automatic analysis triggered from application list.")
        
        resume_file = None
        job_description_file = None

        try:
            # Construct file path from application's resume URL
            relative_path = application.resume_url.lstrip('/').replace('media/', '', 1)
            file_path = os.path.join(settings.MEDIA_ROOT, relative_path)
            
            # This part simulates opening the file from the filesystem.
            # In a production environment, you might use Django's storage API.
            if os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    resume_file = SimpleUploadedFile(
                        name=os.path.basename(file_path),
                        content=f.read(),
                        content_type='application/pdf'
                    )
                logging.info(f"Using resume file from application ID: {application.id}.")
            else:
                messages.error(request, "Resume file not found for the selected application.")
                logging.error(f"Resume file not found at path: {file_path}")
            
            # Get the job description file object from the database
            job_description_doc = JobDescriptionDocument.objects.get(pk=job_description_id)
            job_description_file = job_description_doc.file
        
            if resume_file and job_description_file:
                job_role = job_description_doc.title
                min_years_required = application.experience if application.experience is not None else 0
                
                experience_type = 'range'
                max_years = min_years_required + 5
                
                llm_response = services.analyze_resume_with_llm(
                    resume_file_obj=resume_file,
                    job_description_file_obj=job_description_file,
                    job_role=job_role,
                    experience_type=experience_type,
                    min_years=min_years_required,
                    max_years=max_years,
                )
                
                if llm_response and not llm_response.get("error"):
                    analysis_result = llm_response
                    
                    # --- NEW CODE TO SAVE ANALYSIS TO DATABASE ---
                    try:
                        # Extract and serialize complex fields
                        scoring_matrix_json = json.dumps(analysis_result.get('scoring_matrix', []))
                        bench_recommendation_json = json.dumps(analysis_result.get('bench_recommendation', {}))
                        alternative_role_recommendations_json = json.dumps(analysis_result.get('alternative_role_recommendations', []))
                        automated_recruiter_insights_json = json.dumps(analysis_result.get('automated_recruiter_insights', {}))
                        technical_prowess_json = json.dumps(analysis_result.get('analysis_summary', {}).get('technical_prowess', {}))
                        project_impact_json = json.dumps(analysis_result.get('analysis_summary', {}).get('project_impact', []))
                        education_certifications_json = json.dumps(analysis_result.get('analysis_summary', {}).get('education_certifications', {}))
                        
                        # Create and save a new CandidateAnalysis record
                        CandidateAnalysis.objects.create(
                            full_name=analysis_result.get('full_name'),
                            job_role=job_role,
                            phone_no=analysis_result.get('contact_number'),
                            hiring_recommendation=analysis_result.get('hiring_recommendation'),
                            suggested_salary_range=analysis_result.get('suggested_salary_range'),
                            interview_questions=json.dumps(analysis_result.get('interview_questions', [])),
                            analysis_summary=json.dumps(analysis_result.get('analysis_summary')),
                            experience_match=analysis_result.get('experience_match'),
                            overall_experience=analysis_result.get('overall_experience'),
                            current_company_name=analysis_result.get('current_company_name'),
                            current_company_address=analysis_result.get('current_company_address'),
                            fitment_verdict=analysis_result.get('fitment_verdict'),
                            aggregate_score=analysis_result.get('aggregate_score'),
                            strategic_alignment=analysis_result.get('candidate_fitment_analysis', {}).get('strategic_alignment'),
                            quantifiable_impact=analysis_result.get('candidate_fitment_analysis', {}).get('quantifiable_impact'),
                            potential_gaps_risks=analysis_result.get('candidate_fitment_analysis', {}).get('potential_gaps_risks'),
                            comparable_experience=analysis_result.get('candidate_fitment_analysis', {}).get('comparable_experience_analysis'),
                            scoring_matrix_json=scoring_matrix_json,
                            bench_recommendation_json=bench_recommendation_json,
                            alternative_role_recommendations_json=alternative_role_recommendations_json,
                            automated_recruiter_insights_json=automated_recruiter_insights_json,
                            candidate_overview=analysis_result.get('analysis_summary', {}).get('candidate_overview'),
                            technical_prowess_json=technical_prowess_json,
                            project_impact_json=project_impact_json,
                            education_certifications_json=education_certifications_json,
                            overall_rating_summary=analysis_result.get('analysis_summary', {}).get('overall_rating_summary'),
                            conclusion_summary=analysis_result.get('analysis_summary', {}).get('conclusion_summary'),
                            resume_file_path=application.resume_url,
                            analysis_type='Basic'
                        )
                        logging.info("Candidate analysis saved to database.")
                        messages.success(request, f"AI analysis completed and saved for {analysis_result.get('full_name', 'the candidate')}.")
                    except Exception as db_e:
                        logging.error(f"Error saving analysis to database: {db_e}", exc_info=True)
                        messages.error(request, "AI analysis completed, but there was an error saving the results.")
                    # --- END OF NEW CODE ---
                else:
                    error_message = llm_response.get("error", "AI analysis failed.")
                    messages.error(request, error_message)

        except JobDescriptionDocument.DoesNotExist:
            messages.error(request, "Selected job description not found.")
            logging.error(f"Job description with ID {job_description_id} not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred during analysis: {e}", exc_info=True)
            messages.error(request, f"An unexpected error occurred: {e}")
            
    # Handle GET or POST for initial form display or manual submission
    form_initial_data = {}
    if application:
        form_initial_data = {
            'job_role': application.subject if application.subject else '',
            'min_years_required': application.experience if application.experience is not None else 0,
            'application_id': application.id,
        }
        
    form = ResumeUploadForm(request.POST or None, request.FILES or None, initial=form_initial_data)

    if request.method == 'POST' and form.is_valid():
        pass

    context = {
        'form': form,
        'analysis_result': analysis_result,
        'resume_url': resume_url,
        'job_description_documents': job_description_documents,
    }

    return render(request, 'basic_resume_analysis.html', context)




def advance_resume_analysis_view(request):
    """
    This Django view handles the AI-based resume analysis process.
    It can be triggered automatically from a URL or manually via a form.

    Args:
        request: The Django HttpRequest object.

    Returns:
        A Django HttpResponse that renders the analysis results page.
    """
    analysis_result = None
    resume_url = None
    job_description_documents = JobDescriptionDocument.objects.all()

    # Get application ID and job description ID from URL query parameters
    application_id = request.GET.get('application_id')
    job_description_id = request.GET.get('job_description_id')
    
    application = None
    if application_id:
        try:
            application = Application.objects.get(id=application_id)
            if application.resume_url:
                resume_url = application.resume_url
        except Application.DoesNotExist:
            messages.error(request, "Application not found.")
            logging.error(f"Application with ID {application_id} not found.")

    # Check if analysis should be performed automatically
    if application_id and job_description_id:
        logging.info("Automatic analysis triggered from application list.")
        
        resume_file = None
        job_description_file = None

        try:
            # Construct file path from application's resume URL
            relative_path = application.resume_url.lstrip('/').replace('media/', '', 1)
            file_path = os.path.join(settings.MEDIA_ROOT, relative_path)
            
            # This part simulates opening the file from the filesystem.
            # In a production environment, you might use Django's storage API.
            if os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    resume_file = SimpleUploadedFile(
                        name=os.path.basename(file_path),
                        content=f.read(),
                        content_type='application/pdf'
                    )
                logging.info(f"Using resume file from application ID: {application.id}.")
            else:
                messages.error(request, "Resume file not found for the selected application.")
                logging.error(f"Resume file not found at path: {file_path}")
            
            # Get the job description file object from the database
            job_description_doc = JobDescriptionDocument.objects.get(pk=job_description_id)
            job_description_file = job_description_doc.file
            
            if resume_file and job_description_file:
                # Extract job role from the JobDescriptionDocument object
                job_role = job_description_doc.title
                min_years_required = application.experience if application.experience is not None else 0
                
                # Assign default values for the missing arguments as in your code
                experience_type = 'range'
                max_years = min_years_required + 5  # Set a reasonable upper range
                
                # Call the LLM service
                llm_response = services.analyze_resume_with_llm(
                    resume_file_obj=resume_file,
                    job_description_file_obj=job_description_file,
                    job_role=job_role,
                    experience_type=experience_type,
                    min_years=min_years_required,
                    max_years=max_years,
                )
                
                if llm_response and not llm_response.get("error"):
                    analysis_result = llm_response
                    print(analysis_result)
                    
                    # --- NEW CODE TO SAVE ANALYSIS TO DATABASE ---
                    try:
                        # Extract and serialize complex fields
                        scoring_matrix_json = json.dumps(analysis_result.get('scoring_matrix', []))
                        bench_recommendation_json = json.dumps(analysis_result.get('bench_recommendation', {}))
                        alternative_role_recommendations_json = json.dumps(analysis_result.get('alternative_role_recommendations', []))
                        automated_recruiter_insights_json = json.dumps(analysis_result.get('automated_recruiter_insights', {}))
                        technical_prowess_json = json.dumps(analysis_result.get('analysis_summary', {}).get('technical_prowess', {}))
                        project_impact_json = json.dumps(analysis_result.get('analysis_summary', {}).get('project_impact', []))
                        education_certifications_json = json.dumps(analysis_result.get('analysis_summary', {}).get('education_certifications', {}))
                        
                        # Create and save a new CandidateAnalysis record
                        CandidateAnalysis.objects.create(
                            full_name=analysis_result.get('full_name'),
                            job_role=job_role, # Use the job_role extracted from job_description_doc
                            phone_no=analysis_result.get('contact_number'),
                            hiring_recommendation=analysis_result.get('hiring_recommendation'),
                            suggested_salary_range=analysis_result.get('suggested_salary_range'),
                            interview_questions=json.dumps(analysis_result.get('interview_questions', [])),
                            analysis_summary=json.dumps(analysis_result.get('analysis_summary')),
                            experience_match=analysis_result.get('experience_match'),
                            overall_experience=analysis_result.get('overall_experience'),
                            current_company_name=analysis_result.get('current_company_name'),
                            current_company_address=analysis_result.get('current_company_address'),
                            fitment_verdict=analysis_result.get('fitment_verdict'),
                            aggregate_score=analysis_result.get('aggregate_score'),
                            strategic_alignment=analysis_result.get('candidate_fitment_analysis', {}).get('strategic_alignment'),
                            quantifiable_impact=analysis_result.get('candidate_fitment_analysis', {}).get('quantifiable_impact'),
                            potential_gaps_risks=analysis_result.get('candidate_fitment_analysis', {}).get('potential_gaps_risks'),
                            comparable_experience=analysis_result.get('candidate_fitment_analysis', {}).get('comparable_experience_analysis'),
                            scoring_matrix_json=scoring_matrix_json,
                            bench_recommendation_json=bench_recommendation_json,
                            alternative_role_recommendations_json=alternative_role_recommendations_json,
                            automated_recruiter_insights_json=automated_recruiter_insights_json,
                            candidate_overview=analysis_result.get('analysis_summary', {}).get('candidate_overview'),
                            technical_prowess_json=technical_prowess_json,
                            project_impact_json=project_impact_json,
                            education_certifications_json=education_certifications_json,
                            overall_rating_summary=analysis_result.get('analysis_summary', {}).get('overall_rating_summary'),
                            conclusion_summary=analysis_result.get('analysis_summary', {}).get('conclusion_summary'),
                            resume_file_path=application.resume_url,
                            analysis_type='Advance' 
                        )
                        logging.info("Candidate analysis saved to database.")
                        messages.success(request, f"AI analysis completed and saved for {analysis_result.get('full_name', 'the candidate')}.")
                    except Exception as db_e:
                        logging.error(f"Error saving analysis to database: {db_e}", exc_info=True)
                        messages.error(request, "AI analysis completed, but there was an error saving the results.")
                    # --- END OF NEW CODE ---
                else:
                    error_message = llm_response.get("error", "AI analysis failed.")
                    messages.error(request, error_message)

        except JobDescriptionDocument.DoesNotExist:
            messages.error(request, "Selected job description not found.")
            logging.error(f"Job description with ID {job_description_id} not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred during analysis: {e}", exc_info=True)
            messages.error(request, f"An unexpected error occurred: {e}")
            
    # Handle GET or POST for initial form display or manual submission
    form_initial_data = {}
    if application:
        form_initial_data = {
            'job_role': application.subject if application.subject else '',
            'min_years_required': application.experience if application.experience is not None else 0,
            'application_id': application.id,
        }
    
    # The form is no longer needed for automatic analysis but is kept for context
    form = ResumeUploadForm(request.POST or None, request.FILES or None, initial=form_initial_data)

    if request.method == 'POST' and form.is_valid():
        pass # This section is now vestigial but can be left in for future use.

    context = {
        'form': form,
        'analysis_result': analysis_result,
        'resume_url': resume_url,
        'job_description_documents': job_description_documents,
    }

    return render(request, 'advance_resume_analysis.html', context)

def calendar_scheduler(request):
    return render(request,'calendar_scheduler.html')

#################### Start Send Email Process ###########################


def send_configured_email(user, subject, message_body, recipient_list):
    """
    A helper function to send an email using the dynamic
    EmailConfiguration from the database tied to a specific user.
    """
    try:
        # Get or create the email configuration for the logged-in user.
        config, created = EmailConfiguration.objects.get_or_create(user=user)
        
        from_email = config.email_from if config.email_from else config.email_host_user
        
        connection = config.get_connection()
        email = EmailMessage(
            subject=subject,
            body=message_body,
            from_email=from_email,
            to=recipient_list,
            connection=connection
        )
        
        email.send(fail_silently=False)
        return True
    except EmailConfiguration.DoesNotExist:
        print("Email configuration does not exist for this user.")
        return False
    except Exception as e:
        print(f"Failed to send email: {e}")
        return False

# --- Views ---

@login_required
def configure_email(request):
    """
    Handles the display and submission of the email configuration form.
    For GET requests, it fetches the existing configuration.
    For POST requests, it updates or creates a new configuration.
    """
    # Get or create the email configuration for the current user.
    config, created = EmailConfiguration.objects.get_or_create(user=request.user)

    if request.method == 'POST':
        # Get data from the POST request
        email_host = request.POST.get('email_host')
        email_port = request.POST.get('email_port')
        imap_host = request.POST.get('imap_host')
        imap_port = request.POST.get('imap_port')
        email_host_user = request.POST.get('email_host_user')
        email_host_password = request.POST.get('email_host_password')
        email_from = request.POST.get('email_from')
        security_protocol = request.POST.get('security_protocol')

        # Update the model instance with the new data
        config.email_host = email_host
        config.email_port = email_port
        config.imap_host = imap_host
        config.imap_port = imap_port
        config.email_host_user = email_host_user
        config.email_host_password = email_host_password
        config.email_from = email_from
        
        # Set TLS and SSL based on the radio button selection
        config.email_use_tls = (security_protocol == 'tls')
        config.email_use_ssl = (security_protocol == 'ssl')
        
        try:
            # Save the updated model instance to the database
            config.save()
            messages.success(request, "Email configuration saved successfully!")
            # Redirect to the same page to show the updated data
            return redirect('configure_email')
        except Exception as e:
            messages.error(request, f"Error saving configuration: {e}")
            return redirect('configure_email')

    # For GET requests, render the template with the current configuration
    return render(request, 'email_dashboard.html', {'config': config})


@login_required
def send_job_description(request):
    """
    Handles sending a job description to a candidate via email and logs the sent email.
    """
    job_descriptions = JobDescriptionDocument.objects.all()

    if request.method == 'POST':
        recipient_emails_string = request.POST.get('recipient_emails')
        job_description_id = request.POST.get('job_description')
        email_body = request.POST.get('email_body')
        
        if recipient_emails_string:
            recipient_list = [email.strip() for email in recipient_emails_string.replace(';', ',').split(',') if email.strip()]
        else:
            recipient_list = []

        if not recipient_list or not job_description_id or not email_body:
            messages.error(request, "All fields are required. Please provide at least one recipient email.")
            return render(request, 'send_job_description.html', {'job_descriptions': job_descriptions})

        try:
            job_description = get_object_or_404(JobDescriptionDocument, pk=job_description_id)
            file_name = os.path.basename(job_description.file.name)
            pdf_url = request.build_absolute_uri(f'/media/job_descriptions/{file_name}')
            updated_email_body = f"{email_body}\n\nYou can view the full job description here: {pdf_url}"
            subject = f"Job Description: {job_description.title}"
            
            success = send_configured_email(request.user, subject, updated_email_body, recipient_list)

            if success:
                # Log the sent email
                SentEmail.objects.create(
                    user=request.user,
                    recipient_emails=", ".join(recipient_list),
                    subject=subject,
                    body=updated_email_body
                )
                messages.success(request, "Email sent successfully!")
                return redirect('email_dashboard')
            else:
                messages.error(request, "Failed to send email. Please check your configuration.")
                return render(request, 'send_job_description.html', {'job_descriptions': job_descriptions})

        except Exception as e:
            messages.error(request, f"An error occurred: {e}")
            return render(request, 'send_job_description.html', {'job_descriptions': job_descriptions})
    
    context = {'job_descriptions': job_descriptions}
    return render(request, 'send_job_description.html', context)


@login_required
def sent_emails(request):
    """
    Displays a list of emails sent by the current user.
    """
    sent_emails_list = SentEmail.objects.filter(user_id=request.user.id).order_by('-sent_at')
    context = {'sent_emails_list': sent_emails_list}
    return render(request, 'sent_emails.html', context)

@login_required
def inbox(request):
    """
    Fetches emails from the user's IMAP server and returns them as a JSON response.
    """
    try:
        # Retrieve the user's email configuration from the database
        config = EmailConfiguration.objects.get(user_id=request.user.id)
    except EmailConfiguration.DoesNotExist:
        return JsonResponse({'error_message': 'Email configuration not found. Please configure your email settings.'}, status=404)

    imap_server = None
    try:
        # Establish a connection to the IMAP server
        imap_server = imaplib.IMAP4_SSL(config.imap_host)
        imap_server.login(config.email_host_user, config.email_host_password)
        imap_server.select('INBOX')

        # Search for all email messages in the inbox
        status, messages_data = imap_server.search(None, 'ALL')
        email_ids = messages_data[0].split()
        
        # Get the 10 most recent email IDs
        latest_emails_ids = email_ids[-10:] if len(email_ids) > 10 else email_ids
        
        inbox_emails = []
        for e_id in reversed(latest_emails_ids):
            # Fetch the entire email message
            status, data = imap_server.fetch(e_id, '(RFC822)')
            raw_email = data[0][1]
            msg = email.message_from_bytes(raw_email)

            # Parse email details
            subject = msg.get('Subject')
            from_address = msg.get('From')
            date = msg.get('Date')
            
            # Get the email body
            email_body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get_content_disposition())

                    if content_type == 'text/plain' and 'attachment' not in content_disposition:
                        email_body = part.get_payload(decode=True).decode()
                        break
                else:
                    email_body = msg.get_payload(decode=True).decode()

            # Append the email details to the list
            inbox_emails.append({
                'from': from_address,
                'subject': subject,
                'date': date,
                'body': email_body,
            })

        return JsonResponse({'inbox_emails': inbox_emails})

    except Exception as e:
        # Return a JSON response with an error message
        return JsonResponse({'error_message': f'Failed to connect to IMAP server: {e}'}, status=500)
    finally:
        # Ensure the connection is closed
        if imap_server:
            imap_server.logout()


@login_required
def save_draft(request):
    """
    Saves a new email draft or updates an existing one.
    """
    if request.method == 'POST':
        draft_id = request.POST.get('draft_id')
        subject = request.POST.get('subject', '')
        body = request.POST.get('body', '')
        recipient_emails = request.POST.get('recipient_emails', '')
        
        try:
            if draft_id:
                draft = get_object_or_404(DraftEmail, pk=draft_id, user_id=request.user.id)
                draft.subject = subject
                draft.body = body
                draft.recipient_emails = recipient_emails
                draft.save()
                messages.success(request, "Draft updated successfully!")
            else:
                DraftEmail.objects.create(
                    user=request.user,
                    subject=subject,
                    body=body,
                    recipient_emails=recipient_emails
                )
                messages.success(request, "Draft saved successfully!")
        except Exception as e:
            messages.error(request, f"Failed to save draft: {e}")
            
    return redirect('list_drafts')

@login_required
def list_drafts(request):
    """
    Displays a list of email drafts for the current user.
    """
    
    drafts = DraftEmail.objects.filter(user_id=request.user.id).order_by('-last_modified')
    context = {'drafts': drafts}
    return render(request, 'list_drafts.html', context)

@login_required
def edit_draft(request, draft_id):
    """
    Renders the form to edit a specific draft.
    """
    draft = get_object_or_404(DraftEmail, pk=draft_id, user_id=request.user.id)
    context = {'draft': draft}
    return render(request, 'edit_draft.html', context)

@login_required
def send_draft(request, draft_id):
    """
    Sends an email from a saved draft and deletes the draft upon success.
    """
    draft = get_object_or_404(DraftEmail, pk=draft_id, user_id=request.user.id)
    recipient_list = [email.strip() for email in draft.recipient_emails.replace(';', ',').split(',') if email.strip()]
    
    if not recipient_list:
        messages.error(request, "Cannot send email. No recipient specified in the draft.")
        return redirect('edit_draft', draft_id=draft.id)
    
    success = send_configured_email(request.user, draft.subject, draft.body, recipient_list)
    
    if success:
        # Log the sent email
        SentEmail.objects.create(
            user=request.user,
            recipient_emails=draft.recipient_emails,
            subject=draft.subject,
            body=draft.body
        )
        draft.delete()
        messages.success(request, "Email sent successfully and draft deleted!")
        return redirect('email_dashboard')
    else:
        messages.error(request, "Failed to send email. Please check your configuration.")
        return redirect('email_dashboard')

@login_required
def delete_draft(request, draft_id):
    """
    Deletes a specific draft.
    """
    if request.method == 'POST':
        draft = get_object_or_404(DraftEmail, pk=draft_id, user_id=request.user.id)
        draft.delete()
        messages.success(request, "Draft deleted successfully!")
    return redirect('list_drafts')


@login_required
def success_page(request):
    """
    A simple success page view.
    """
    return HttpResponse("Email sent successfully!")

@login_required
def get_job_description_content(request, job_id):
    """
    Fetches job description content for a given ID.
    This view is for the AJAX call in the frontend.
    """
    try:
        job_description = get_object_or_404(JobDescriptionDocument, pk=job_id)
        return JsonResponse({'description': job_description.description})
    except Exception as e:
        return JsonResponse({'error': f'Job description not found: {e}'}, status=404)
    

@login_required
def email_dashboard(request):
    """
    A unified view for the email dashboard, displaying inbox, sent, drafts,
    and the form to send a new email.
    """
    context = {}
    error_message = None

    # --- Fetch all necessary data ---

    # Get email configuration
    try:
        config = EmailConfiguration.objects.get(user_id=request.user.id)
        context['config'] = config
    except EmailConfiguration.DoesNotExist:
        context['config'] = None
        error_message = "Email configuration not found. Please configure your email settings."

    # Fetch sent emails
    sent_emails_list = SentEmail.objects.filter(user_id=request.user.id).order_by('-sent_at')
    context['sent_emails_list'] = sent_emails_list

    # Fetch drafts
    drafts = DraftEmail.objects.filter(user_id=request.user.id).order_by('-last_modified')
    context['drafts'] = drafts

    # Fetch job descriptions for the "send new email" form
    job_descriptions = JobDescriptionDocument.objects.all()
    context['job_descriptions'] = job_descriptions
    
    # --- Inbox logic: only include replies to platform-sent emails ---
    inbox_emails = []
    if not error_message:  # Only try to fetch if configuration exists
        imap_server = None
        try:
            # Get list of all recipients you have ever emailed from this platform
            platform_recipients = set()
            for sent in sent_emails_list:
                if sent.recipient_emails:
                    platform_recipients.update([email.strip() for email in sent.recipient_emails.split(";")])

            # Establish a connection to the IMAP server
            imap_server = imaplib.IMAP4_SSL(config.imap_host)
            imap_server.login(config.email_host_user, config.email_host_password)
            imap_server.select('INBOX')

            # Search for all email messages in the inbox
            status, messages_data = imap_server.search(None, 'ALL')
            email_ids = messages_data[0].split()

            # Get the 20 most recent email IDs (increase/decrease as needed)
            latest_emails_ids = email_ids[-20:] if len(email_ids) > 20 else email_ids

            for e_id in reversed(latest_emails_ids):
                # Fetch the entire email message
                status, data = imap_server.fetch(e_id, '(RFC822)')
                raw_email = data[0][1]
                msg = email.message_from_bytes(raw_email)

                # Parse email details
                subject = msg.get('Subject', '')
                from_address = msg.get('From', '')
                date = msg.get('Date', '')

                # Only include if the "From" matches one of our platform recipients
                if not any(recipient in from_address for recipient in platform_recipients):
                    continue

                # Get the email body
                email_body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        content_disposition = str(part.get_content_disposition())

                        if content_type == 'text/plain' and 'attachment' not in content_disposition:
                            try:
                                email_body = part.get_payload(decode=True).decode()
                            except:
                                email_body = part.get_payload(decode=True).decode(errors="ignore")
                            break
                else:
                    try:
                        email_body = msg.get_payload(decode=True).decode()
                    except:
                        email_body = msg.get_payload(decode=True).decode(errors="ignore")

                # Append the email details to the list
                inbox_emails.append({
                    'from': from_address,
                    'subject': subject,
                    'date': date,
                    'body': email_body,
                })

        except Exception as e:
            error_message = f'Failed to connect to IMAP server: {e}'
        finally:
            if imap_server:
                imap_server.logout()

    context['inbox_emails'] = inbox_emails
    context['error_message'] = error_message
    
    return render(request, 'email_dashboard.html', context)




###################### Start User create #################

@user_passes_test(is_admin_or_superadmin, login_url='/signin/')
def create_user(request):
    """
    Handles the form submission for creating a new user.
    """
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        role = request.POST.get('role', 'user')  # Default to 'user' if not specified

        # Ensure that only superadmins can create admin users
        if request.user.role == 'admin' and role != 'user':
            messages.error(request, "You do not have permission to create this type of user.")
            return redirect('admin_dashboard')
        
        # Ensure that the role selected is valid based on the user's permissions
        if request.user.role == 'superadmin' and role not in ['admin', 'user']:
            messages.error(request, "Invalid role selected.")
            return redirect('superadmin_dashboard')
            
        try:
            # Create a new user with the provided details
            new_user = User.objects.create_user(username=username, password=password, role=role)

            # --- This crucial block of code links the new user to their creator ---
            if request.user.role == 'admin':
                new_user.created_by = request.user
            
            # --- This new line ensures that any new admin has staff status enabled ---
            if role == 'admin' or role == 'superadmin':
                new_user.is_staff = True

            new_user.save()
            # ---------------------------------------------------------------------

            messages.success(request, f"User '{new_user.username}' created successfully with role '{new_user.get_role_display()}'.")
        except Exception as e:
            messages.error(request, f"Error creating user: {e}")
            
        if request.user.role == 'superadmin':
            return redirect('superadmin_dashboard')
        else:
            return redirect('admin_dashboard')

    # If the request is not a POST, redirect to the appropriate dashboard
    if request.user.role == 'superadmin':
        return redirect('superadmin_dashboard')
    else:
        return redirect('admin_dashboard')
 

    
# Helper function to check if the user is a superuser or admin
def is_admin_or_superuser(user):
    return user.is_superuser or user.groups.filter(name='Admins').exists()

@login_required
@user_passes_test(is_admin_or_superuser)
def toggle_user_status(request, user_id):
    if request.method == 'POST':
        target_user = get_object_or_404(User, id=user_id)

        # Prevent an admin from disabling their own account
        if target_user == request.user:
            messages.error(request, "You cannot change the status of your own account.")
            return redirect('admin_dashboard')

        target_user.is_active = not target_user.is_active
        target_user.save()

        if target_user.is_active:
            messages.success(request, f"User '{target_user.username}' has been enabled.")
        else:
            messages.success(request, f"User '{target_user.username}' has been disabled.")

    return redirect('admin_dashboard')

def set_user_password(request, user_id):
    """
    Sets a manual password for a specific user.
    """
    if request.method == 'POST':
        try:
            user = User.objects.get(id=user_id)
            new_password = request.POST.get('new_password')
            
            if not new_password:
                messages.error(request, "Password cannot be empty.")
                return redirect('admin_dashboard')

            # Set the new password for the user
            user.password = make_password(new_password)
            user.save()
            
            # Add a success message to be displayed in the pop-up
            messages.success(request, f"Password for '{user.username}' has been successfully set.")
            
        except User.DoesNotExist:
            messages.error(request, "User not found.")
            
        return redirect('admin_dashboard')

    return redirect('admin_dashboard')


#################  Job Posting ################
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium import webdriver
from selenium.webdriver.common.by import By



@login_required
def post_jobs(request):
    """
    A Django view to retrieve job descriptions from the database
    and post them to LinkedIn.
    """
    try:
        # Fetch all job descriptions from the database.
        # This is a simplified query; you can add filters (e.g., jobs not yet posted).
        jobs_to_post = JobDescriptionDocument.objects.all()

        # Convert the Django queryset to a list of dictionaries
        jobs_data = list(jobs_to_post.values(
            'title', 'company_name', 'job_level', 'employment_type',
            'country', 'state', 'city', 'overview', 'responsibilities',
            'required_skills', 'preferred_skills', 'education_experience',
            'benefits'
        ))

        # Check if any jobs were found
        if not jobs_data:
            return HttpResponse("No job descriptions found to post.", status=200)

        # Call the Selenium function to post the jobs
        post_jobs_to_linkedin(
            jobs_data=jobs_data,
            username="rahulsuthar7280@gmail.com",
            password="rahul@7280"
        )

        return HttpResponse(f"Successfully started the job posting process for {len(jobs_data)} jobs.", status=200)
    
    except Exception as e:
        # In a real-world scenario, you would log this error and return a more user-friendly message.
        return HttpResponse(f"An error occurred while trying to post the jobs: {str(e)}", status=500)



########################## Career page ###################
# This view will only handle the initial GET request to render the page

# def settings_careerpage(request):
#     """
#     Handles displaying and processing all three forms:
#     1. CareerPage (New Job Posting)
#     2. CompanyInfo
#     3. Job Application (Apply_career)
#     """
    
#     # Handle POST requests
#     if request.method == 'POST':
#         form_type = request.POST.get('form_type')

#         if form_type == 'job_application':
#             # Handle Job Application form submission
#             career_id = request.POST.get('career')
#             career_instance = CareerPage.objects.get(id=career_id)

#             Apply_career.objects.create(
#                 # Link to the user if they are logged in, otherwise it's null
#                 user=request.user if request.user.is_authenticated else None,
#                 career=career_instance,
#                 first_name=request.POST.get('first_name'),
#                 last_name=request.POST.get('last_name'),
#                 email=request.POST.get('email'),
#                 phone=request.POST.get('phone'),
#                 experience=request.POST.get('experience'),
#                 current_ctc=request.POST.get('current_ctc'),
#                 expected_ctc=request.POST.get('expected_ctc'),
#                 qualification=request.POST.get('qualification'),
#                 notice_period=request.POST.get('notice_period'),
#                 resume=request.FILES.get('resume'),
#                 cover_letter=request.FILES.get('cover_letter'),
#                 linkedin_url=request.POST.get('linkedin_url'),
#             )
#             messages.success(request, "Your job application has been submitted successfully!")
#             return redirect('settings_careerpage')

#         elif form_type == 'company_info':
#             # Handle Company Info form submission
#             company_info, created = CompanyInfo.objects.get_or_create(id=1) # Get or create the single instance

#             company_info.company_name = request.POST.get('company_name')
#             company_info.address = request.POST.get('address')
#             company_info.phone_number = request.POST.get('phone_number')
#             company_info.email = request.POST.get('email')
#             company_info.about_us_url = request.POST.get('about_us_url')
#             company_info.contact_us_url = request.POST.get('contact_us_url')
#             company_info.our_services_url = request.POST.get('our_services_url')
#             company_info.privacy_policy_url = request.POST.get('privacy_policy_url')
#             company_info.terms_and_conditions_url = request.POST.get('terms_and_conditions_url')
#             company_info.twitter_url = request.POST.get('twitter_url')
#             company_info.facebook_url = request.POST.get('facebook_url')
#             company_info.youtube_url = request.POST.get('youtube_url')
#             company_info.linkedin_url = request.POST.get('linkedin_url')
#             company_info.save()

#             messages.success(request, "Company information has been updated successfully!")
#             return redirect('settings_careerpage')

#         elif form_type == 'career_page':
#             # Handle new Job Posting form submission
#             new_job = CareerPage.objects.create(
#                 title=request.POST.get('title'),
#                 company=request.POST.get('company'),
#                 company_logo=request.FILES.get('company_logo'),
#                 location=request.POST.get('location'),
#                 job_type=request.POST.get('job_type'),
#                 experience=request.POST.get('experience'),
#                 salary=request.POST.get('salary'),
#                 is_active=request.POST.get('is_active', False) == 'on',
#                 description=request.POST.get('description'),
#                 about_company=request.POST.get('about_company'),
#                 skills=request.POST.get('skills'),
#                 benefits=request.POST.get('benefits'),
#                 application_link=request.POST.get('application_link'),
#                 responsibilities=request.POST.get('responsibilities'),
#                 qualifications=request.POST.get('qualifications'),
#                 date_line=request.POST.get('date_line'),
#             )
#             messages.success(request, f"New job posting for '{new_job.title}' created successfully!")
#             return redirect('settings_careerpage')

#     # Handle GET requests (to display the forms)
#     context = {}
#     return render(request, 'settings_careerpage.html', context)


########################### Start setting_careerpage ##################
# @login_required
# --- ASSUMED MODELS IMPORTS ---
# NOTE: Ensure these models are defined in your models.py and linked to the User model.
# The model imports you provided
try:
    # Assuming models are in the same app or correctly configured
    from .models import JobApplicationFormSettings, CompanyInfo, ThemeSettings, CareerPage, Category
except ImportError:
    # Raise a clear error if the models are missing, as the function depends on them
    raise Exception("Required models (JobApplicationFormSettings, CompanyInfo, ThemeSettings, CareerPage, Category) must be defined and imported.")
# ------------------------------

# Define common icon choices (if not defined on the Category model)
# Note: You use a mix of 'fas fa-X' and 'fa-X' in your code. Using 'fas fa-X' here for consistency.
ICON_CHOICES = [
    ('fas fa-code', 'Code'),
    ('fas fa-chart-line', 'Chart'),
    ('fas fa-headset', 'Headset'),
    ('fas fa-palette', 'Palette'),
    ('fas fa-star', 'Star'),
    ('fas fa-link', 'Link'),
    ('fas fa-calendar', 'Calendar'),
    ('fas fa-map-marker-alt', 'Location'),
    ('fas fa-hashtag', 'Hashtag'),
    ('fas fa-edit', 'Edit')
]


@login_required
def settings_careerpage(request):
    """
    Handles all configuration settings for the user's career page, including 
    Job Application Form fields, Company Info, Theme Colors, and Job/Category management.
    """
    
    # -----------------------------------------------------------
    # 1. Retrieve or Create all necessary user settings objects
    # -----------------------------------------------------------
    user = request.user
    
    # Retrieves or creates settings objects linked to the current user
    form_settings, _ = JobApplicationFormSettings.objects.get_or_create(user=user)
    company_info, _ = CompanyInfo.objects.get_or_create(user=user)
    theme_settings, _ = ThemeSettings.objects.get_or_create(user=user)
    
    # Resolve the base URL once to handle redirects with fragments correctly
    try:
        # Assumes the URL is named 'settings_careerpage' in urls.py
        base_url = reverse('settings_careerpage') 
    except Exception:
        # Fallback if the URL name is not set up
        base_url = '/settings_careerpage/' 

    # -----------------------------------------------------------
    # 2. Handle POST requests for different form types
    # -----------------------------------------------------------
    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        # --- Job Application Form Logic ---
        if form_type == 'job_application':
            try:
                # Update standard fields
                form_settings.first_name_enabled = 'first_name_enabled' in request.POST
                form_settings.last_name_enabled = 'last_name_enabled' in request.POST
                form_settings.email_enabled = 'email_enabled' in request.POST
                form_settings.phone_enabled = 'phone_enabled' in request.POST
                form_settings.experience_enabled = 'experience_enabled' in request.POST
                form_settings.current_ctc_enabled = 'current_ctc_enabled' in request.POST
                form_settings.expected_ctc_enabled = 'expected_ctc_enabled' in request.POST
                form_settings.notice_period_enabled = 'notice_period_enabled' in request.POST
                form_settings.qualification_enabled = 'qualification_enabled' in request.POST
                form_settings.linkedin_url_enabled = 'linkedin_url_enabled' in request.POST
                form_settings.resume_enabled = 'resume_enabled' in request.POST
                form_settings.cover_letter_enabled = 'cover_letter_enabled' in request.POST

                # Handling Additional Fields
                additional_fields_data = []
                # Use a more robust way to iterate over posted additional fields
                field_ids = sorted([int(k.split('_')[-1]) for k in request.POST.keys() if k.startswith('additional_field_name_')])
                
                for field_id in field_ids:
                    field_name = request.POST.get(f'additional_field_name_{field_id}', '').strip()
                    is_enabled = f'additional_field_enabled_{field_id}' in request.POST
                    
                    if is_enabled and field_name:
                        additional_fields_data.append({
                            'name': field_name,
                            'type': request.POST.get(f'additional_field_type_{field_id}', 'text'),
                            'icon_class': request.POST.get(f'additional_field_icon_{field_id}', 'fas fa-edit'),
                            'id': field_id
                        })
                
                form_settings.additional_fields = additional_fields_data
                
                form_settings.save()
                messages.success(request, "Job application form settings updated successfully!")
                return redirect(base_url + '#job-app-form') 

            except Exception as e:
                messages.error(request, f"Error updating application form settings: {e}")
                return redirect(base_url + '#job-app-form')

        # --- Company Info Logic ---
        elif form_type == 'company_info':
            try:
                company_info.company_name = request.POST.get('company_name')
                
                # Handle company logo upload
                if 'company_logo' in request.FILES:
                    company_info.company_logo = request.FILES.get('company_logo')
                    
                company_info.address = request.POST.get('address')
                company_info.phone_number = request.POST.get('phone_number')
                company_info.email = request.POST.get('email')
                company_info.about_us_url = request.POST.get('about_us_url')
                company_info.contact_us_url = request.POST.get('contact_us_url')
                company_info.our_services_url = request.POST.get('our_services_url')
                company_info.privacy_policy_url = request.POST.get('privacy_policy_url')
                company_info.terms_and_conditions_url = request.POST.get('terms_and_conditions_url')
                company_info.twitter_url = request.POST.get('twitter_url')
                company_info.facebook_url = request.POST.get('facebook_url')
                company_info.youtube_url = request.POST.get('youtube_url')
                company_info.linkedin_url = request.POST.get('linkedin_url')
                company_info.save()
                messages.success(request, "Company information has been updated successfully!")
                return redirect(base_url + '#company-info-form')

            except Exception as e:
                messages.error(request, f"Error updating company information: {e}")
                return redirect(base_url + '#company-info-form')


        # --- Theme Settings Logic ---
        elif form_type == 'theme_settings':
            try:
                # 1. Background/UI Colors
                theme_settings.theme_primary_color = request.POST.get('primary_color', '#1e3a8a')
                theme_settings.theme_secondary_color = request.POST.get('secondary_color', '#eef2f6')
                theme_settings.theme_background_color = request.POST.get('background_color', '#f7f9fc')
                
                # 2. Text Colors (The new 3 fields)
                theme_settings.theme_primary_color_text = request.POST.get('primary_color_text', '#ffffff')
                theme_settings.theme_secondary_color_text = request.POST.get('secondary_color_text', '#1f2937')
                theme_settings.theme_background_color_text = request.POST.get('background_color_text', '#1f2937')
                
                theme_settings.save()
                messages.success(request, "Theme settings updated successfully!")
                return redirect(base_url + '#theme-form')
                
            except Exception as e:
                messages.error(request, f"Error updating theme settings: {e}")
                return redirect(base_url + '#theme-form')
                
        # --- Category Logic (Create) ---
        elif form_type == 'category':
            try:
                category_name = request.POST.get('category_name')
                icon_class = request.POST.get('icon_class')
                
                if not category_name:
                    messages.error(request, "Category name cannot be empty.")
                else:
                    Category.objects.create(
                        user=user, 
                        name=category_name,
                        icon_class=icon_class
                    )
                    messages.success(request, f"Category '{category_name}' added successfully!")
                return redirect(base_url + '#category-form')

            except IntegrityError:
                messages.error(request, f"A category named '{category_name}' already exists.")
                return redirect(base_url + '#category-form')
            except Exception as e:
                messages.error(request, f"Error creating category: {e}")
                return redirect(base_url + '#category-form')

        # --- Career Page (New Job Posting) Logic ---
        elif form_type == 'career_page':
            try:
                category_name = request.POST.get('category')
                
                # Find the category object for validation/existence check
                # Even though CareerPage.category is a CharField, this ensures a valid category name is used.
                get_object_or_404(Category, user=user, name=category_name) 

                new_job = CareerPage.objects.create(
                    user=user,
                    title=request.POST.get('title'),
                    company=request.POST.get('company'),
                    location=request.POST.get('location'),
                    job_type=request.POST.get('job_type'),
                    # FIX: Save the string name to match the CareerPage model's CharField
                    category=category_name, 
                    experience=request.POST.get('experience'),
                    salary=request.POST.get('salary'),
                    is_active=request.POST.get('is_active') == 'on',
                    description=request.POST.get('description'),
                    responsibilities=request.POST.get('responsibilities'),
                    qualifications=request.POST.get('qualifications'),
                    benefits=request.POST.get('benefits'),
                    date_line=parse_date(request.POST.get('date_line')) if request.POST.get('date_line') else None,
                )
                messages.success(request, f"New job posting for '{new_job.title}' created successfully!")
                return redirect(base_url + '#career-page-form') 
                
            except Exception as e:
                messages.error(request, f"Error creating job posting: {e}")
                return redirect(base_url + '#career-page-form')

    # -----------------------------------------------------------
    # 3. Handle GET request / Context preparation
    # -----------------------------------------------------------
    
    # Fetch all jobs for the user
    jobs = CareerPage.objects.filter(user=user).order_by('-id')

    # Fetch all categories for the user
    categories = Category.objects.filter(user=user)
    print("categories",categories)

    # Use ORM aggregation to efficiently get the job counts per category
    category_counts_queryset = CareerPage.objects.filter(user=user).values('category').annotate(count=Count('category'))
    
    # Convert the queryset into a dictionary for easy lookup
    vacancy_counts = {item['category']: item['count'] for item in category_counts_queryset}
    print("vacancy_counts",vacancy_counts)

    # Add the count to each category object
    categories_with_counts = []
    for category in categories:
        count = vacancy_counts.get(category.name, 0)
        # Use an attribute to store the count for the template
        category.vacancy_count = count
        categories_with_counts.append(category)

    # Use the defined ICON_CHOICES if the model doesn't have them
    icon_choices = getattr(Category, 'ICON_CHOICES', ICON_CHOICES)
    
    context = {
        'form_settings': form_settings,
        'company_info': company_info,
        'theme_settings': theme_settings,
        'jobs': jobs,
        'categories': categories_with_counts,
        'icon_choices': icon_choices, 
    }
    
    return render(request, 'settings_careerpage.html', context)


@login_required
def edit_job(request, job_id):
    job = get_object_or_404(CareerPage, pk=job_id, user=request.user)
    
    if request.method == 'POST':
        # Job Edit Logic
        job.title = request.POST.get('title')
        job.company = request.POST.get('company')
        if 'company_logo' in request.FILES:
            job.company_logo = request.FILES.get('company_logo')
        job.location = request.POST.get('location')
        job.job_type = request.POST.get('job_type')
        job.experience = request.POST.get('experience')
        job.salary = request.POST.get('salary')
        job.is_active = request.POST.get('is_active', False) == 'on'
        job.description = request.POST.get('description')
        job.about_company = request.POST.get('about_company')
        job.skills = request.POST.get('skills')
        job.benefits = request.POST.get('benefits')
        job.application_link = request.POST.get('application_link')
        job.responsibilities = request.POST.get('responsibilities')
        job.qualifications = request.POST.get('qualifications')
        job.date_line = request.POST.get('date_line')
        job.save()
        messages.success(request, f"Job posting for '{job.title}' updated successfully!")
        return redirect('settings_careerpage')

    context = {
        'job': job,
    }
    return render(request, 'edit_job.html', context)


@login_required
def delete_job(request, job_id):
    job = get_object_or_404(CareerPage, pk=job_id, user=request.user)
    job_title = job.title # Save name before deleting
    job.delete()
    messages.success(request, f"Job posting for '{job_title}' deleted successfully!")
    return redirect('settings_careerpage')


@login_required
def edit_category(request, category_id):
    category = get_object_or_404(Category, pk=category_id, user=request.user)
    
    # Resolve base URL here to redirect back with the fragment
    base_url = reverse('settings_careerpage')

    if request.method == 'POST':
        category.name = request.POST.get('category_name')
        category.vacancy_count = request.POST.get('vacancy_count', 0)
        category.icon_class = request.POST.get('icon_class')
        category.save()
        messages.success(request, f"Category updated to '{category.name}' successfully!")
        # FIX: Append fragment to the resolved URL
        return redirect(base_url + '#category-form') 

    # Pass the ICON_CHOICES to the template for editing
    icon_choices = Category.ICON_CHOICES
    
    context = {
        'category': category,
        'icon_choices': icon_choices, 
    }
    return render(request, 'edit_category.html', context)


@login_required
def delete_category(request, category_id):
    category = get_object_or_404(Category, pk=category_id, user=request.user)
    category_name = category.name # Save name before deleting
    category.delete()
    messages.success(request, f"Category '{category_name}' deleted successfully!")
    return redirect('settings_careerpage')

################### close setting_careerpage #######################

def list_careers(request):
    """
    Displays the list of all career pages.
    """
    all_jobs = list(CareerPage.objects.all().values())
    
    jobs_json = json.dumps(all_jobs, default=str)
    
    locations = CareerPage.objects.values_list('location', flat=True).distinct()
    job_types = CareerPage.objects.values_list('job_type', flat=True).distinct()
    
    context = {
        'jobs_data_json': jobs_json,
        'locations': locations,
        'job_types': job_types,
    }
    
    return render(request, 'career_portal.html', context)

# Helper function to get theme settings, reducing redundancy
def get_theme_settings(user=None):
    """Fetches ThemeSettings, prioritizing the given user or the first entry."""
    # Importing inside the function allows for the model to be outside views.py
    from .models import ThemeSettings 

    # 1. Try to get the settings for the specific user (e.g., the job poster)
    if user:
        try:
            # Note: The model instance will have the full field names
            return ThemeSettings.objects.get(user=user)
        except ThemeSettings.DoesNotExist:
            pass # Fall through to default

    # 2. Fallback: Get the first available ThemeSettings (site-wide default)
    # Assuming the first object, if it exists, is the site-wide default when user is None
    try:
        return ThemeSettings.objects.first()
    except Exception:
        pass # Fall through to hardcoded default

    # 3. Hardcoded default if no settings exist
    class DefaultTheme:
        # --- Background Colors (Matching the model field names) ---
        theme_primary_color = '#1e3a8a'
        theme_secondary_color = '#eef2f6'
        theme_background_color = '#f7f9fc'
        
        # --- Text Colors (NEW) ---
        theme_primary_color_text = '#ffffff'
        theme_secondary_color_text = '#333333'
        theme_background_color_text = '#333333'
        
    return DefaultTheme()

@require_http_methods(["GET", "POST"])
def job_detail(request, job_id):
    """
    Renders the job detail page and handles job application submissions, including dynamic additional fields.
    """
    try:
        job = get_object_or_404(CareerPage, pk=job_id)
        form_settings, created = JobApplicationFormSettings.objects.get_or_create(user=job.user)
        theme_settings, created = ThemeSettings.objects.get_or_create(user=job.user)
        company_info, created = CompanyInfo.objects.get_or_create(user=job.user)
    except (CareerPage.DoesNotExist, JobApplicationFormSettings.DoesNotExist, ThemeSettings.DoesNotExist, CompanyInfo.DoesNotExist):
        return render(request, 'error_page.html', {'message': 'Requested content not found.'}, status=404)

    if request.method == 'POST':
        # Server-side validation for 'experience'
        experience_value = request.POST.get('experience', None)
        if experience_value and not experience_value.isdigit():
            messages.error(request, 'Total Experience must be a positive integer.')
            return redirect('job_detail', job_id=job_id)

        try:
            # Prepare a dictionary to store additional field data
            additional_data = {}
            # Iterate through the fields enabled in the settings and get their values
            for i, field in enumerate(form_settings.additional_fields):
                field_name = field.get('name')
                # The name in the form is "additional_field_1", "additional_field_2", etc.
                form_key = f'additional_field_{i + 1}'
                additional_data[field_name] = request.POST.get(form_key, '')

            # Create a new Apply_career instance with data from the POST request
            Apply_career.objects.create(
                career=job,
                first_name=request.POST.get('first_name', ''),
                last_name=request.POST.get('last_name', ''),
                email=request.POST.get('email', ''),
                phone=request.POST.get('phone_number', ''),
                experience=experience_value,
                current_ctc=request.POST.get('current_ctc', ''),
                expected_ctc=request.POST.get('expected_ctc', ''),
                qualification=request.POST.get('qualification', ''),
                notice_period=request.POST.get('notice_period', ''),
                linkedin_url=request.POST.get('linkedin_url', ''),
                resume=request.FILES.get('resume'),
                cover_letter=request.FILES.get('cover_letter', None),
                additional_data=additional_data  # Save the dynamically collected data here
            )
            messages.success(request, 'Your application has been submitted successfully!')
            # Redirect to the main career page
            return redirect('career_mainpage')
        except Exception as e:
            messages.error(request, f'An error occurred: {e}. Please try again.')
            return redirect('job_detail', job_id=job_id)

    context = {
        'job': job,
        'form_settings': form_settings,
        'theme_settings': theme_settings,
        'company_info': company_info,
        'company_name': company_info
    }
    
    return render(request, 'job_detail.html', context)

def career_mainpage(request):
    """
    Renders the main career page with dynamic data.
    """
    
    # Use your actual model calls to fetch data from the database
    categories = Category.objects.all()
    
    # Filter jobs based on their is_active status
    featured_jobs = CareerPage.objects.filter(job_type='Featured', is_active=True)
    full_time_jobs = CareerPage.objects.filter(job_type='Full Time', is_active=True)
    part_time_jobs = CareerPage.objects.filter(job_type='Part Time', is_active=True)
    
    # Fetch the company information
    try:
        company_info = CompanyInfo.objects.first()
    except CompanyInfo.DoesNotExist:
        company_info = None

    # Get the site-wide theme settings
    theme_settings = ThemeSettings.objects.first()

    # Pass the data to the template
    context = {
        'categories': categories,
        'featured_jobs': featured_jobs,
        'full_time_jobs': full_time_jobs,
        'part_time_jobs': part_time_jobs,
        'company_info': company_info,
        'theme_settings': theme_settings,  # Pass the entire object to the template
    }
    return render(request, 'career_mainpage.html', context)

def add_job_listing(request):
    """
    Handles form submission to add new job listings without forms.py.
    """
    categories = Category.objects.all() # Fetch categories for the select dropdown

    if request.method == 'POST':
        # Create a new Job object from the POST data
        try:
            new_job = Job(
                title=request.POST.get('job_title'),
                location=request.POST.get('job_location'),
                job_type=request.POST.get('job_type'),
                salary_range=request.POST.get('salary_range'),
                date_line=request.POST.get('date_line'),
                company_logo=request.FILES.get('company_logo'),
                category=Category.objects.get(id=request.POST.get('category'))
            )
            new_job.save()
            return redirect('career_mainpage')
        except Exception as e:
            # Handle potential errors, e.g., if a category ID is invalid
            print(f"Error saving job: {e}")
            pass # Or handle the error more gracefully

    context = {
        'categories': categories
    }
    return render(request, 'add_job.html', context)





@csrf_exempt
@require_http_methods(["POST"])
def create_job(request):
    """
    Handles the creation of a new job posting via a POST request from JavaScript.
    """
    try:
        data = json.loads(request.body)
        skills_str = json.dumps(data.get('skills', []))
        
        # Determine salary or stipend based on job type
        job_type = data.get('jobType')
        salary_value = data.get('salary')
        
        job = CareerPage.objects.create(
            title=data.get('title'),
            company=data.get('company'),
            location=data.get('location'),
            job_type=job_type,
            experience=data.get('experienceLevel'),
            salary=salary_value,
            description=data.get('description'),
            responsibilities=data.get('responsibilities'),
            qualifications=data.get('qualifications'),
            benefits=data.get('benefits'),
            skills=skills_str,
            is_active=data.get('is_active', True)
        )
        return JsonResponse({'id': job.id, 'success': True, 'message': 'Job created successfully.'}, status=201)
    except (json.JSONDecodeError, KeyError) as e:
        return JsonResponse({'success': False, 'message': 'Invalid data provided.'}, status=400)

@csrf_exempt
@require_http_methods(["PUT"])
def update_job(request, job_id):
    """
    Handles the update of an existing job posting via a PUT request from JavaScript.
    """
    try:
        job = get_object_or_404(CareerPage, pk=job_id)
        data = json.loads(request.body)
        
        job.title = data.get('title', job.title)
        job.company = data.get('company', job.company)
        job.location = data.get('location', job.location)
        job.job_type = data.get('jobType', job.job_type)
        job.experience = data.get('experienceLevel', job.experience)
        job.salary = data.get('salary', job.salary)
        job.description = data.get('description', job.description)
        job.responsibilities = data.get('responsibilities', job.responsibilities)
        job.qualifications = data.get('qualifications', job.qualifications)
        job.benefits = data.get('benefits', job.benefits)
        
        # Handle skills list, converting back to JSON string
        skills_list = data.get('skills', [])
        job.skills = json.dumps(skills_list)
        
        job.save()
        
        return JsonResponse({'id': job.id, 'success': True, 'message': 'Job updated successfully.'})
    except (json.JSONDecodeError, KeyError) as e:
        return JsonResponse({'success': False, 'message': 'Invalid data provided.'}, status=400)

@csrf_exempt
@require_http_methods(["DELETE"])
def delete_job(request, job_id):
    """
    Handles the deletion of a job posting via a DELETE request from JavaScript.
    """
    job = get_object_or_404(CareerPage, pk=job_id)
    job.delete()
    return JsonResponse({'success': True, 'message': 'Job deleted successfully.'})

@csrf_exempt
@require_http_methods(["POST"])
def toggle_job_status(request, job_id):
    """
    Toggles the active status of a job posting.
    """
    try:
        data = json.loads(request.body)
        is_active = data.get('is_active')
        job = get_object_or_404(CareerPage, pk=job_id)
        job.is_active = is_active
        job.save()
        return JsonResponse({'id': job.id, 'is_active': job.is_active, 'success': True, 'message': 'Status updated.'})
    except (json.JSONDecodeError, KeyError) as e:
        return JsonResponse({'success': False, 'message': 'Invalid data provided.'}, status=400)

@csrf_exempt
def apply_for_job(request, job_id):
    if request.method == 'POST':
        career_page = get_object_or_404(CareerPage, pk=job_id)

        application = Apply_career(
            career=career_page,
            first_name=request.POST.get('first_name'),
            last_name=request.POST.get('last_name'),
            email=request.POST.get('email'),
            phone=request.POST.get('phone'),
            experience=request.POST.get('experience'),
            current_ctc=request.POST.get('current_ctc'),
            expected_ctc=request.POST.get('expected_ctc'),
            qualification=request.POST.get('qualification'),
            notice_period=request.POST.get('notice_period'),
            linkedin_url=request.POST.get('linkedin_url'),
            resume=request.FILES.get('resume'),
            cover_letter=request.FILES.get('cover_letter')
        )
        application.save()
        messages.success(request, 'Your application has been submitted successfully!')
        return redirect('job_detail', job_id=job_id)

    return redirect('job_detail', job_id=job_id)


########################################################

@require_http_methods(["POST", "DELETE"])
def manage_job(request, job_id):
    """
    Handles editing and deleting a job posting.
    """
    job = get_object_or_404(CareerPage, pk=job_id)

    if request.method == 'POST':
        # Handle Edit functionality
        # Update job fields from POST data
        job.title = request.POST.get('job-title')
        job.company = request.POST.get('company-name')
        job.location = request.POST.get('specific-location')
        job.job_type = request.POST.get('job-type')
        job.experience = request.POST.get('experience-level')
        job.description = request.POST.get('job-description')
        job.about_company = request.POST.get('about-company')
        job.skills = request.POST.get('skills')
        job.benefits = request.POST.get('benefits')
        job.application_link = request.POST.get('application-link')
        
        # Determine salary or stipend
        if job.job_type == 'Internship':
            salary_value = request.POST.get('stipend-amount') or 'N/A'
        else:
            min_salary = request.POST.get('min-salary')
            max_salary = request.POST.get('max-salary')
            if min_salary and max_salary:
                salary_value = f'${min_salary} - ${max_salary}'
            elif min_salary:
                salary_value = f'From ${min_salary}'
            elif max_salary:
                salary_value = f'Up to ${max_salary}'
            else:
                salary_value = 'N/A'
        job.salary = salary_value

        # Handle logo update
        if 'company-logo' in request.FILES:
            job.company_logo = request.FILES['company-logo']
        
        job.save()
        return JsonResponse({'message': 'Job updated successfully!'}, status=200)

    elif request.method == 'DELETE':
        # Handle Delete functionality
        job.delete()
        return JsonResponse({'message': 'Job deleted successfully!'}, status=200)
    
    return JsonResponse({'error': 'Invalid request method'}, status=400)
# The other views (toggle_status, share_career, career_detail) remain unchanged
def toggle_status(request, pk):
    career_page = get_object_or_404(CareerPage, pk=pk)
    career_page.is_active = not career_page.is_active
    career_page.save()
    return redirect('career_portal')

def share_career(request, pk):
    career_page = get_object_or_404(CareerPage, pk=pk)
    share_url = request.build_absolute_uri(reverse('career_detail', args=[pk]))
    return redirect('list_careers')

def career_detail(request, pk):
    career_page = get_object_or_404(CareerPage, pk=pk)
    
    if request.method == 'POST':
        # Manually retrieve form data from the POST request
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        linkedin_url = request.POST.get('linkedin_url')
        resume_file = request.FILES.get('resume')
        cover_letter_file = request.FILES.get('cover_letter')
        
        # Perform a basic check for required fields
        if not all([first_name, last_name, email, resume_file]):
            # You might want to add error messages to the context here
            return render(request, 'career_detail.html', {'career': career_page, 'error': 'Please fill out all required fields.'})
        
        # Handle file uploads and create the file path
        fs = FileSystemStorage()
        resume_filename = fs.save(resume_file.name, resume_file)
        
        # Save cover letter if provided
        cover_letter_filename = None
        if cover_letter_file:
            cover_letter_filename = fs.save(cover_letter_file.name, cover_letter_file)

        # Create and save a new Application instance, include the user field
        application = Apply_career.objects.create(
            user=request.user,  # This is the authenticated user
            career=career_page,
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone,
            resume=fs.url(resume_filename),
            cover_letter=fs.url(cover_letter_filename) if cover_letter_filename else None,
            linkedin_url=linkedin_url
        )
        
        return redirect('career_detail', pk=career_page.pk)
    
    context = {
        'career': career_page,
    }
    return render(request, 'career_detail.html', context)

def application_success(request):
    """
    A simple view for a success page after form submission.
    """
    return render(request, 'application_success.html')



##################### Create Folder ##################
from django.db.models import Count
@login_required # Protect this view from unauthenticated access
def file_manager_view(request, folder_id=None):
    """
    Displays either all folders or the contents of a specific folder for the logged-in user.
    """
    selected_folder = None
    
    # Use annotate to get the count of documents for each folder, filtered by user
    all_folders = Folder.objects.filter(user=request.user).annotate(
        document_count=Count('documents')
    ).order_by('name')

    if folder_id is not None:
        # State 2: Folder-Browse State
        selected_folder = get_object_or_404(all_folders, pk=folder_id, user=request.user) # Filter by user
        documents_to_display = selected_folder.documents.all()
        return render(request, 'file_manager.html', {
            'selected_folder': selected_folder,
            'documents_to_display': documents_to_display,
            'all_folders': all_folders
        })
    else:
        # State 1: Home State
        return render(request, 'file_manager.html', {
            'all_folders': all_folders,
            'selected_folder': None
        })

# --- Folder Management Views ---

@login_required
def create_folder_view(request):
    """
    Handles POST request to create a new folder for the logged-in user.
    """
    if request.method == 'POST':
        folder_name = request.POST.get('folder_name')
        if folder_name:
            # Check for existing folder with the same name for the current user
            if Folder.objects.filter(user=request.user, name=folder_name).exists():
                messages.warning(request, f"Folder '{folder_name}' already exists.")
            else:
                Folder.objects.create(user=request.user, name=folder_name)
                messages.success(request, f"Folder '{folder_name}' created successfully.")
    return redirect('file_manager')

@login_required
def delete_folder_view(request, folder_id):
    """
    Handles POST request to delete a folder and its documents for the logged-in user.
    """
    if request.method == 'POST':
        # Get the folder, making sure it belongs to the current user
        folder = get_object_or_404(Folder, pk=folder_id, user=request.user)
        folder.delete()
        messages.success(request, f"Folder '{folder.name}' and all its contents have been deleted.")
    return redirect('file_manager')

@login_required
def edit_folder_view(request, folder_id):
    """
    Handles POST request to rename a folder for the logged-in user.
    """
    if request.method == 'POST':
        new_name = request.POST.get('new_name')
        if new_name:
            # Get the folder, ensuring it belongs to the current user
            folder = get_object_or_404(Folder, pk=folder_id, user=request.user)
            old_name = folder.name
            folder.name = new_name
            folder.save()
            messages.success(request, f"Folder '{old_name}' renamed to '{new_name}'.")
    return redirect('file_manager')

# --- Document Management Views ---

@login_required
def delete_document_view(request, document_id):
    """
    Handles POST request to delete a document for the logged-in user.
    """
    if request.method == 'POST':
        # Get the document, ensuring it belongs to the current user
        document = get_object_or_404(Document, pk=document_id, user=request.user)
        folder_id = document.folder.pk
        document.delete()
        messages.success(request, f"Document '{document.file.name}' has been deleted.")
    return redirect('browse_folder', folder_id=folder_id)

@login_required
def upload_file_view(request):
    """
    Handles POST request to upload files to a folder for the logged-in user.
    """
    if request.method == 'POST':
        folder_id = request.POST.get('folder_id')
        files = request.FILES.getlist('file')
        try:
            # Get the folder, ensuring it belongs to the current user
            folder = Folder.objects.get(pk=folder_id, user=request.user)
            for file in files:
                # Create the document and associate it with the user and folder
                Document.objects.create(folder=folder, file=file, user=request.user)
            messages.success(request, f"{len(files)} file(s) uploaded successfully to '{folder.name}'.")
        except Folder.DoesNotExist:
            messages.error(request, "Folder not found.")
    return redirect('file_manager')

# --- AI Auto-Sort Views ---
# (The 'extract_text_from_file' and 'get_job_role_with_llm' functions remain unchanged as they don't directly handle user-specific data, but are called from a user-authenticated view)

def extract_text_from_file(file):
    """
    Extracts text from various file types.
    """
    file_content = ""
    # file_name = file.name.lower()
    file_name = str(file).lower()

    
    # Check if the file is a text file
    if file_name.endswith('.txt'):
        try:
            file_content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            file_content = file.read().decode('latin-1')
    
    # Check if the file is a PDF
    elif file_name.endswith('.pdf'):
        try:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                file_content += page.extract_text() or ''
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            file_content = ""

    # Check if the file is a DOCX
    elif file_name.endswith('.docx'):
        try:
            doc = docx.Document(file)
            for para in doc.paragraphs:
                file_content += para.text + '\n'
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            file_content = ""

    return file_content

def get_job_role_with_llm(file_content):
    """
    Identifies a job role from file content using the Google Gemini LLM.
    """
    if not file_content:
        return "Uncategorized"

    prompt = (
        "Analyze the following resume content and identify the primary job role. "
        "Return only a single, capitalized job title or a general category. "
        "If the role is not clearly defined, return 'Uncategorized'.\n\n"
        "Example output: 'Python Developer', 'Data Scientist', 'Project Manager', 'Marketing Specialist', 'Uncategorized'\n\n"
        "Resume Content:\n"
        f"{file_content[:2000]}..."  # Truncate content to save tokens
    )

    try:
        model = genai.GenerativeModel('gemini-1.5-flash') # Use an appropriate Gemini model
        response = model.generate_content(prompt)
        
        llm_output = response.text.strip()
        
        # Basic cleaning of the LLM output to ensure it is a valid folder name
        cleaned_name = re.sub(r'[^a-zA-Z0-9\s]', '', llm_output).strip()
        if not cleaned_name or cleaned_name.lower() == "uncategorized":
            return "Uncategorized"
        
        return cleaned_name
        
    except Exception as e:
        print(f"Gemini API request failed: {e}")
        return "Uncategorized"

@login_required
def ai_auto_sort_view(request):
    """
    Handles POST request to upload and automatically sort files using an LLM for the logged-in user.
    """
    if request.method == 'POST':
        files = request.FILES.getlist('files')
        if not files:
            messages.error(request, "No files were selected for sorting.")
            return redirect('file_manager')

        successful_sorts = 0
        for file in files:
            try:
                # 1. Read the file content
                file_content = extract_text_from_file(file)
                if not file_content:
                    messages.warning(request, f"Could not extract text from file '{file.name}'. Skipping.")
                    continue

                # 2. Use LLM to determine the folder name
                folder_name = get_job_role_with_llm(file_content)
                
                # 3. Get or create the folder, filtered by the current user
                folder, created = Folder.objects.get_or_create(user=request.user, name=folder_name)

                # 4. Save the file to the determined folder and link it to the user
                Document.objects.create(folder=folder, file=file, user=request.user)
                successful_sorts += 1

            except Exception as e:
                messages.error(request, f"Failed to process file {file.name}: {e}")
                continue

        messages.success(request, f"{successful_sorts} file(s) have been successfully sorted by AI.")
    return redirect('file_manager')
######################## Close Folder ########################

